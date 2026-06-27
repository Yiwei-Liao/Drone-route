from __future__ import annotations

import csv
import json
import math
import re
import shutil
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageOps

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "docs" / "report"
FIG_DIR = REPORT_DIR / "figures"
MATLAB_DIR = REPORT_DIR / "matlab"
MATLAB_FIG_DIR = REPORT_DIR / "figures_matlab"
REPORT_PATH = REPORT_DIR / "泰山低空巡检无人机通信链路可靠性影响机制与优化策略研究.docx"
NOTES_PATH = REPORT_DIR / "report_notes.md"
REVISION_NOTES_PATH = REPORT_DIR / "revision_notes.md"


TITLE = "泰山低空巡检无人机通信链路可靠性影响机制与优化策略研究"
SUBTITLE = "基于现有航线、杆塔空间数据、三维仿真沙盘与文献传播模型的技术预评估报告"


def ensure_dirs() -> None:
    for path in [REPORT_DIR, FIG_DIR, MATLAB_DIR, MATLAB_FIG_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def as_float(value: str | None) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except ValueError:
        return None


def summarize_data() -> dict[str, object]:
    routes = read_csv_rows(ROOT / "data" / "processed" / "routes.csv")
    waypoints = read_csv_rows(ROOT / "data" / "processed" / "route_waypoints.csv")
    towers = read_csv_rows(ROOT / "data" / "processed" / "towers.csv")
    stations = read_csv_rows(ROOT / "data" / "processed" / "base_stations.csv")
    tasks = read_csv_rows(ROOT / "data" / "processed" / "flight_tasks.csv")
    matches = read_csv_rows(ROOT / "data" / "processed" / "route_tower_matches.csv")
    manifest_path = ROOT / "data" / "processed" / "manifest.json"
    manifest = {}
    if manifest_path.exists():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

    valid_towers = [
        row
        for row in towers
        if as_float(row.get("longitude")) is not None and as_float(row.get("latitude")) is not None
    ]
    route_stats = []
    for row in routes:
        route_stats.append(
            {
                "name": row.get("route_name", ""),
                "kml_file": row.get("kml_file", ""),
                "waypoint_count": int(float(row.get("waypoint_count") or 0)),
                "length_m": float(row.get("total_length") or 0),
                "min_height": float(row.get("min_height") or 0),
                "max_height": float(row.get("max_height") or 0),
                "avg_height": float(row.get("avg_height") or 0),
                "nearest": "",
            }
        )
    match_by_route = {row.get("route_id"): row for row in matches}
    for item, route in zip(route_stats, routes):
        match = match_by_route.get(route.get("route_id"))
        if match:
            item["nearest"] = f"{match.get('line_name', '')}{match.get('tower_no', '')}"
            item["min_distance_to_tower"] = float(match.get("min_distance_to_tower") or 0)
            item["match_confidence"] = match.get("match_confidence", "")

    return {
        "routes": routes,
        "waypoints": waypoints,
        "towers": towers,
        "valid_towers": valid_towers,
        "stations": stations,
        "tasks": tasks,
        "matches": matches,
        "manifest": manifest,
        "route_stats": route_stats,
        "source_files": manifest.get("input_files", []),
        "processed_files": manifest.get("processed_files", []),
    }


def font_path() -> str:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return ""


FONT_PATH = font_path()


def make_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        try:
            return ImageFont.truetype(FONT_PATH, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def draw_text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "#1b2630",
) -> None:
    lines = text.split("\n")
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + (len(lines) - 1) * 8
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for line, w, h in zip(lines, widths, line_heights):
        x = box[0] + (box[2] - box[0] - w) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + 8


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    outline: str = "#35546a",
    radius: int = 16,
    width: int = 3,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = "#295f85",
    width: int = 5,
) -> None:
    draw.line([start, end], fill=fill, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (
        end[0] - size * math.cos(ang - math.pi / 6),
        end[1] - size * math.sin(ang - math.pi / 6),
    )
    p2 = (
        end[0] - size * math.cos(ang + math.pi / 6),
        end[1] - size * math.sin(ang + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=fill)


def antialias_canvas(width: int = 1600, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw, int]:
    scale = 2
    image = Image.new("RGB", (width * scale, height * scale), "white")
    draw = ImageDraw.Draw(image)
    return image, draw, scale


def finish_canvas(image: Image.Image, path: Path, width: int = 1600, height: int = 900) -> None:
    image = image.resize((width, height), Image.Resampling.LANCZOS)
    image.save(path, quality=95, dpi=(220, 220))


def plot_axes(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    xlim: tuple[float, float],
    ylim: tuple[float, float],
    xlabel: str,
    ylabel: str,
    title: str,
    scale: int,
) -> tuple[callable, callable]:
    left, top, right, bottom = [v * scale for v in box]
    font = make_font(22 * scale)
    small = make_font(18 * scale)
    title_font = make_font(28 * scale)
    grid = "#d9dde1"
    axis = "#354052"
    draw.rectangle((left, top, right, bottom), fill="#ffffff", outline=axis, width=2 * scale)
    for i in range(1, 6):
        x = left + (right - left) * i / 6
        draw.line((x, top, x, bottom), fill=grid, width=1 * scale)
    for i in range(1, 5):
        y = top + (bottom - top) * i / 5
        draw.line((left, y, right, y), fill=grid, width=1 * scale)

    def tx(x: float) -> int:
        return int(left + (x - xlim[0]) / (xlim[1] - xlim[0]) * (right - left))

    def ty(y: float) -> int:
        return int(bottom - (y - ylim[0]) / (ylim[1] - ylim[0]) * (bottom - top))

    draw.text((left, top - 50 * scale), title, font=title_font, fill="#14213d")
    draw.text(((left + right) / 2 - 70 * scale, bottom + 36 * scale), xlabel, font=font, fill=axis)
    draw.text((left - 92 * scale, top + 8 * scale), ylabel, font=font, fill=axis)
    for tick in np.linspace(xlim[0], xlim[1], 6):
        x = tx(float(tick))
        draw.line((x, bottom, x, bottom + 8 * scale), fill=axis, width=2 * scale)
        draw.text((x - 22 * scale, bottom + 12 * scale), f"{tick:g}", font=small, fill=axis)
    for tick in np.linspace(ylim[0], ylim[1], 6):
        y = ty(float(tick))
        draw.line((left - 8 * scale, y, left, y), fill=axis, width=2 * scale)
        draw.text((left - 70 * scale, y - 10 * scale), f"{tick:g}", font=small, fill=axis)
    return tx, ty


def make_report_figures(summary: dict[str, object]) -> dict[str, Path]:
    figures: dict[str, Path] = {}

    def copy_or_placeholder(src: Path, dst: Path, title: str) -> Path:
        if src.exists():
            shutil.copyfile(src, dst)
            return dst
        image = Image.new("RGB", (1600, 900), "#f4f7f8")
        draw = ImageDraw.Draw(image)
        draw_text_center(
            draw,
            (100, 320, 1500, 560),
            f"当前未找到 {src.name}\n请运行前端验证脚本后替换本图\n{title}",
            make_font(44),
            "#334155",
        )
        image.save(dst, quality=95, dpi=(220, 220))
        return dst

    web_sources = [
        (
            "web_terrain",
            ROOT / "output" / "terrain-2d-check.png",
            FIG_DIR / "fig_01_web_terrain_imagery_overview.png",
            "地形影像模式：三条 KML 航线与杆塔簇总览",
        ),
        (
            "web_topographic",
            ROOT / "output" / "verified-street-map.png",
            FIG_DIR / "fig_02_web_topographic_route_overview.png",
            "地形影像/拓扑底图模式：巡检线路覆盖范围",
        ),
        (
            "web_3d",
            ROOT / "output" / "current-3d-view.png",
            FIG_DIR / "fig_03_web_3d_route_detail.png",
            "三维模式：局部航线、杆塔与地形关系",
        ),
        (
            "web_basic",
            ROOT / "output" / "map-layer-check-final.png",
            FIG_DIR / "fig_04_web_basic_grid_mode.png",
            "基础网格模式：无外部地理底图时的显示状态",
        ),
    ]
    for key, src, dst, title in web_sources:
        figures[key] = copy_or_placeholder(src, dst, title)

    composite = Image.new("RGB", (1600, 1120), "white")
    draw = ImageDraw.Draw(composite)
    title_font = make_font(32)
    label_font = make_font(23)
    draw.text((55, 30), "网页端多地图模式与巡检线路覆盖对照", font=title_font, fill="#0f2f4a")
    panels = [
        (figures["web_terrain"], (50, 95, 775, 500), "A 地形影像模式：三条 KML 航线空间覆盖"),
        (figures["web_topographic"], (825, 95, 1550, 500), "B 拓扑底图模式：杆塔簇与航线总览"),
        (figures["web_3d"], (50, 590, 775, 995), "C 三维模式：局部航线与地形关系"),
        (figures["web_basic"], (825, 590, 1550, 995), "D 基础网格模式：底图降级与系统状态"),
    ]
    for path, box, label in panels:
        img = Image.open(path).convert("RGB")
        img = ImageOps.fit(img, (box[2] - box[0], box[3] - box[1]), Image.Resampling.LANCZOS)
        composite.paste(img, (box[0], box[1]))
        draw.rectangle(box, outline="#334155", width=3)
        draw.text((box[0], box[3] + 12), label, font=label_font, fill="#334155")
    draw.text(
        (55, 1050),
        "说明：截图来自项目网页端验证输出，地形影像模式覆盖三条 KML 航线簇；三维模式用于局部空间关系复核；基础网格模式用于外部底图不可用时的降级显示说明。",
        font=make_font(21),
        fill="#475569",
    )
    comparison_path = FIG_DIR / "fig_05_web_map_modes_comparison.png"
    composite.save(comparison_path, quality=95, dpi=(220, 220))
    figures["web_comparison"] = comparison_path

    # Fig 06: height/path/fresnel schematic.
    image, draw, s = antialias_canvas()
    W, H = image.size
    font = make_font(24 * s)
    small = make_font(20 * s)
    title = make_font(34 * s)
    draw.text((70 * s, 38 * s), "不同飞行高度下传播路径、山体遮挡和第一菲涅尔区净空示意（模型示意）", font=title, fill="#0f2f4a")
    base_y = 720 * s
    xs = np.linspace(0, 5, 200)
    terrain = 70 + 38 * np.sin(xs * 1.7) + 185 * np.exp(-((xs - 2.8) / 0.55) ** 2) + 55 * np.exp(-((xs - 4.1) / 0.38) ** 2)

    def sx(x: float) -> int:
        return int(130 * s + x / 5 * 1320 * s)

    def sy(h: float) -> int:
        return int(base_y - h / 360 * 560 * s)

    points = [(sx(float(x)), sy(float(h))) for x, h in zip(xs, terrain)]
    poly = [(sx(0), base_y), *points, (sx(5), base_y)]
    draw.polygon(poly, fill="#d9ead3", outline="#719f6a")
    for x, h in zip(xs[::8], terrain[::8]):
        draw.line((sx(float(x)), sy(float(h)), sx(float(x)), base_y), fill="#eef4ec", width=1 * s)
    g = (sx(0.15), sy(95))
    draw.ellipse((g[0] - 14 * s, g[1] - 14 * s, g[0] + 14 * s, g[1] + 14 * s), fill="#275d83")
    draw.text((g[0] - 52 * s, g[1] + 22 * s), "地面站", font=small, fill="#243b53")
    uav_x = 4.65
    heights = [(150, "#bf3f3f", "低高度：遮挡风险高"), (220, "#d39b22", "中高度：净空改善"), (300, "#1d7a47", "高高度：路径损耗增加")]
    for height, color, label in heights:
        u = (sx(uav_x), sy(height))
        draw.line((g[0], g[1], u[0], u[1]), fill=color, width=4 * s)
        mid = ((g[0] + u[0]) // 2, (g[1] + u[1]) // 2)
        if height == 220:
            draw.ellipse((mid[0] - 285 * s, mid[1] - 55 * s, mid[0] + 285 * s, mid[1] + 55 * s), outline="#5aa9e6", width=3 * s)
            draw.text((mid[0] - 225 * s, mid[1] - 90 * s), "第一菲涅尔区 F1，工程常用 0.6F1 净空", font=small, fill="#16609b")
        draw.polygon([(u[0], u[1] - 18 * s), (u[0] - 24 * s, u[1] + 15 * s), (u[0] + 24 * s, u[1] + 15 * s)], fill=color)
        draw.text((u[0] - 270 * s, u[1] - 18 * s), label, font=small, fill=color)
    draw.text((sx(2.35), sy(295)), "山体遮挡会使 LoS 转为 NLoS", font=font, fill="#7a1f1f")
    draw.text((sx(0.3), 805 * s), "注：地形轮廓为模型示意，真实应用需接入 DEM/DSM 数据。", font=small, fill="#475569")
    finish_canvas(image, FIG_DIR / "fig_06_height_path_fresnel.png")
    figures["height"] = FIG_DIR / "fig_06_height_path_fresnel.png"

    # Fig 07: DEM and 0.6F1 clearance.
    image, draw, s = antialias_canvas()
    x = np.linspace(0, 4, 260)
    dem = 92 + 26 * np.sin(x * 2.0) + 118 * np.exp(-((x - 2.25) / 0.42) ** 2) + 35 * np.exp(-((x - 3.2) / 0.25) ** 2)
    h_g, h_u = 80, 235
    h_los = h_g + x / x[-1] * (h_u - h_g)
    lam = 3e8 / (2.4e9)
    D = x[-1] * 1000
    f1 = np.sqrt(np.maximum(0, lam * (x * 1000) * (D - x * 1000) / D))
    boundary = h_los - 0.6 * f1
    tx, ty = plot_axes(
        draw,
        (135, 145, 1460, 710),
        (0, 4),
        (40, 260),
        "链路距离 x (km)",
        "高程/高度 (m)",
        "DEM 地形剖面与 0.6F1 净空判据（模型示意）",
        s,
    )
    pts_dem = [(tx(float(a)), ty(float(b))) for a, b in zip(x, dem)]
    pts_los = [(tx(float(a)), ty(float(b))) for a, b in zip(x, h_los)]
    pts_bd = [(tx(float(a)), ty(float(b))) for a, b in zip(x, boundary)]
    draw.line(pts_dem, fill="#466b2e", width=5 * s)
    draw.line(pts_los, fill="#1f77b4", width=5 * s)
    draw.line(pts_bd, fill="#d9842b", width=5 * s)
    for i in range(0, len(x), 12):
        draw.line((pts_bd[i][0], pts_bd[i][1], pts_los[i][0], pts_los[i][1]), fill="#f2c596", width=1 * s)
    draw.text((1060 * s, 176 * s), "h_LoS(x)", font=make_font(22 * s), fill="#1f77b4")
    draw.text((1020 * s, 240 * s), "h_LoS(x) - 0.6F1(x)", font=make_font(22 * s), fill="#d9842b")
    draw.text((700 * s, 655 * s), "h_DEM(x)", font=make_font(22 * s), fill="#466b2e")
    rounded_box(draw, (1050 * s, 735 * s, 1495 * s, 835 * s), "#f8fafc", "#9aa7b1", radius=12 * s, width=2 * s)
    draw.text((1070 * s, 755 * s), "C(x)=h_LoS-h_DEM-0.6F1\nC(x)>0 表示满足 0.6F1 净空", font=make_font(20 * s), fill="#243b53")
    draw.text((135 * s, 820 * s), "注：本图为 DEM 接入后的判据示意，定量应用需配置项目 DEM/DSM。", font=make_font(20 * s), fill="#475569")
    finish_canvas(image, FIG_DIR / "fig_07_dem_profile_fresnel_clearance.png")
    figures["dem"] = FIG_DIR / "fig_07_dem_profile_fresnel_clearance.png"

    # Fig 08: FSPL curve.
    image, draw, s = antialias_canvas()
    d = np.linspace(0.5, 5.0, 360)
    freqs = [(0.92, "#26734d", "920 MHz"), (2.4, "#1f77b4", "2.4 GHz"), (5.8, "#bf3f3f", "5.8 GHz")]
    all_y = [92.45 + 20 * np.log10(f) + 20 * np.log10(d) for f, _, _ in freqs]
    y_min, y_max = 82, 124
    tx, ty = plot_axes(
        draw,
        (135, 145, 1460, 710),
        (0.5, 5.0),
        (y_min, y_max),
        "链路距离 d (km)",
        "FSPL (dB)",
        "不同频段自由空间路径损耗曲线（理论计算）",
        s,
    )
    x3, x4 = tx(3), tx(4)
    draw.rectangle((x3, ty(y_max), x4, ty(y_min)), fill="#fff3cd")
    draw.text((x3 + 15 * s, ty(121)), "3–4 km 工程敏感距离区间", font=make_font(21 * s), fill="#7a5a00")
    for (freq, color, label), y in zip(freqs, all_y):
        pts = [(tx(float(a)), ty(float(b))) for a, b in zip(d, y)]
        draw.line(pts, fill=color, width=5 * s)
        draw.text((tx(4.65), ty(float(y[-1])) - 20 * s), label, font=make_font(22 * s), fill=color)
    draw.text((160 * s, 770 * s), "L_FSPL = 92.45 + 20log10(f_GHz) + 20log10(d_km)。曲线用于理论链路预算与频段敏感性分析。", font=make_font(20 * s), fill="#475569")
    finish_canvas(image, FIG_DIR / "fig_08_fspl_curve_sensitive_distance.png")
    figures["fspl"] = FIG_DIR / "fig_08_fspl_curve_sensitive_distance.png"

    # Fig 09: EMI influence path.
    image, draw, s = antialias_canvas()
    title_font = make_font(34 * s)
    box_font = make_font(22 * s)
    small_font = make_font(18 * s)
    draw.text((70 * s, 38 * s), "电力设备电磁干扰对无人机通信链路影响路径图（机制示意）", font=title_font, fill="#0f2f4a")
    left_boxes = [
        ((80, 160, 430, 250), "高压导线\n电晕放电"),
        ((80, 300, 430, 390), "绝缘子/金具\n局部放电"),
        ((80, 440, 430, 530), "工频电场/磁场\n近场暴露"),
        ((80, 580, 430, 670), "开关暂态\n金属反射/散射"),
    ]
    mid_boxes = [
        ((610, 210, 980, 330), "前门耦合\nFront-door Coupling\n经天线进入接收机"),
        ((610, 455, 980, 575), "后门耦合\nBack-door Coupling\n经机体线缆/接口进入系统"),
    ]
    right_boxes = [
        ((1180, 145, 1510, 235), "遥控链路\n控制可靠性"),
        ((1180, 260, 1510, 350), "图传/数传链路\n吞吐与时延"),
        ((1180, 375, 1510, 465), "GNSS/北斗\n定位质量"),
        ((1180, 490, 1510, 580), "磁罗盘/飞控\n姿态与导航"),
        ((1180, 605, 1510, 695), "任务载荷\n图像与传感器"),
    ]
    for box, text in left_boxes:
        b = tuple(v * s for v in box)
        rounded_box(draw, b, "#eef6ff", "#5b8fb9", radius=12 * s, width=3 * s)
        draw_text_center(draw, b, text, box_font, "#153e5c")
    for box, text in mid_boxes:
        b = tuple(v * s for v in box)
        rounded_box(draw, b, "#fff7e6", "#d39b22", radius=12 * s, width=3 * s)
        draw_text_center(draw, b, text, small_font, "#5b3a00")
    for box, text in right_boxes:
        b = tuple(v * s for v in box)
        rounded_box(draw, b, "#f1f8f4", "#5a9b68", radius=12 * s, width=3 * s)
        draw_text_center(draw, b, text, box_font, "#1f4d2b")
    for _, _text in left_boxes:
        draw_arrow(draw, (430 * s, 345 * s), (610 * s, 270 * s), "#527da3", width=4 * s)
        draw_arrow(draw, (430 * s, 485 * s), (610 * s, 515 * s), "#a46c1b", width=4 * s)
        break
    for y in [190, 305, 420, 535, 650]:
        draw_arrow(draw, (980 * s, 270 * s), (1180 * s, y * s), "#527da3", width=3 * s)
        draw_arrow(draw, (980 * s, 515 * s), (1180 * s, y * s), "#a46c1b", width=3 * s)
    rounded_box(draw, (470 * s, 700 * s, 1135 * s, 825 * s), "#f8fafc", "#94a3b8", radius=12 * s, width=2 * s)
    draw.text((500 * s, 724 * s), "R_EMI = a1S_dist + a2S_tower + a3S_weather + a4S_equipment", font=make_font(24 * s), fill="#0f2f4a")
    draw.text((500 * s, 768 * s), "该指标用于电磁暴露风险排序，现场强度评估需结合频谱监测和设备抗扰度试验。", font=make_font(20 * s), fill="#475569")
    finish_canvas(image, FIG_DIR / "fig_09_emi_coupling_paths.png")
    figures["emi"] = FIG_DIR / "fig_09_emi_coupling_paths.png"

    # Fig 10: risk workflow.
    image, draw, s = antialias_canvas()
    draw.text((70 * s, 38 * s), "通信链路综合风险分级与优化策略流程图（工程假设）", font=make_font(34 * s), fill="#0f2f4a")
    steps = [
        ((75, 170, 345, 280), "输入数据\n航点/杆塔/基站\n气象/设备参数"),
        ((420, 170, 690, 280), "传播计算\n距离/FSPL\nLoS/Fresnel"),
        ((765, 170, 1035, 280), "EMI 暴露\n近塔距离\n设备/天气因子"),
        ((1110, 170, 1515, 280), "综合风险\nR_total = ΣwiRi\n低/中/高/极高"),
    ]
    for box, text in steps:
        b = tuple(v * s for v in box)
        rounded_box(draw, b, "#eef6ff", "#3d7ca6", radius=14 * s, width=3 * s)
        draw_text_center(draw, b, text, make_font(22 * s), "#153e5c")
    for i in range(len(steps) - 1):
        end_x = steps[i + 1][0][0] * s
        y = 225 * s
        draw_arrow(draw, (steps[i][0][2] * s, y), (end_x, y), "#3d7ca6", width=5 * s)
    controls = [
        ((120, 430, 445, 540), "高度优化\n改善视距，控制路径损耗"),
        ((485, 430, 810, 540), "节点/中继优化\n缩短斜距，绕开遮挡"),
        ((850, 430, 1175, 540), "链路冗余抗干扰\n频段/天线/编码备份"),
        ((1215, 430, 1515, 540), "近塔安全控制\n限距、姿态、失联保护"),
    ]
    for box, text in controls:
        b = tuple(v * s for v in box)
        rounded_box(draw, b, "#f1f8f4", "#5a9b68", radius=14 * s, width=3 * s)
        draw_text_center(draw, b, text, make_font(21 * s), "#1f4d2b")
        draw_arrow(draw, ((box[0] + box[2]) // 2 * s, 430 * s), (1285 * s, 280 * s), "#74a87a", width=3 * s)
    risk_box = (130 * s, 660 * s, 1470 * s, 815 * s)
    rounded_box(draw, risk_box, "#fffdf7", "#d4a72c", radius=12 * s, width=3 * s)
    levels = [
        ("低", "R<0.25：LoS 清晰，0.6F1 净空满足"),
        ("中", "0.25–0.50：净空裕度不足或距离偏大"),
        ("高", "0.50–0.75：LoS 遮挡、3–4 km 区间叠加 EMI"),
        ("极高", "≥0.75：遮挡、距离、EMI 或气象多因素耦合"),
    ]
    x0 = 165 * s
    for i, (lvl, text) in enumerate(levels):
        y = (688 + i * 30) * s
        draw.text((x0, y), f"{lvl}风险", font=make_font(22 * s), fill="#7a5a00")
        draw.text((x0 + 135 * s, y), text, font=make_font(20 * s), fill="#334155")
    finish_canvas(image, FIG_DIR / "fig_10_integrated_risk_workflow.png")
    figures["risk"] = FIG_DIR / "fig_10_integrated_risk_workflow.png"

    matlab_figures = {
        "matlab_fspl": "matlab_fig_01_fspl_curve.png",
        "matlab_fresnel": "matlab_fig_02_fresnel_radius.png",
        "matlab_dem": "matlab_fig_03_dem_fresnel_clearance.png",
        "matlab_height_margin": "matlab_fig_04_height_clearance_margin.png",
        "matlab_risk_heatmap": "matlab_fig_05_distance_height_risk_heatmap.png",
        "matlab_emi_distance": "matlab_fig_06_emi_distance_risk.png",
    }
    for key, filename in matlab_figures.items():
        figures[key] = MATLAB_FIG_DIR / filename

    return figures


def fspl_db(freq_ghz: float, distance_km: float) -> float:
    return 92.45 + 20 * math.log10(freq_ghz) + 20 * math.log10(distance_km)


def set_east_asia_font(run, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)


def set_paragraph_font(paragraph, font_name: str = "宋体", size: float = 12, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_east_asia_font(run, font_name)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12.5, "24415F"),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.text = "泰山低空巡检无人机通信链路可靠性影响机制与优化策略研究"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, "宋体", 9)


def add_title_page(doc: Document, summary: dict[str, object]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_east_asia_font(run, "微软雅黑")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("12395B")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run(SUBTITLE)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("4B5563")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("研究属性", "面向方案论证、模型推导和任务预评估的技术研究报告。"),
        ("数据基础", f"{len(summary['routes'])} 条 KML 航线、{len(summary['waypoints'])} 个航点、{len(summary['valid_towers'])}/{len(summary['towers'])} 个有效杆塔坐标、{len(summary['stations'])} 个基站候选点。"),
        ("核心模型", "自由空间路径损耗、LoS/NLoS、第一菲涅尔区净空、EMI 暴露风险与综合风险指数。"),
        ("验证条件", "当前无 RSSI、SNR、丢包率、图传码率、频谱监测和真实断链位置等观测量，相关数据可用于后续模型校准。"),
        ("生成日期", "2026-06-11"),
    ]
    fill_table(table, rows, widths=[Cm(3.2), Cm(12.6)], header=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("摘要：")
    set_east_asia_font(r, "微软雅黑")
    r.bold = True
    r.font.size = Pt(12)
    r = p.add_run(
        "本报告围绕合同关注的飞行高度影响机制、电力设备电磁干扰影响机制和通信可靠性优化策略展开。"
        "当前项目已具备航线、航点高度、杆塔空间坐标、基站候选点与三维仿真沙盘，可支撑传播机理、敏感区识别和优化策略的技术预评估；"
        "通信质量观测量与 DEM/DSM 数据尚待补充，因此本文采用可解释的风险分量、理论链路预算和工程验证条件组织论证。"
    )
    set_east_asia_font(r, "宋体")
    r.font.size = Pt(12)

    doc.add_page_break()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    add_rich_text(p, text, "宋体", size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width) -> None:
    cell.width = width
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def fill_table(table, rows: list[tuple[str, ...]], widths: list, header: bool = True) -> None:
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx].cells if r_idx == 0 else table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=header and r_idx == 0, size=9.5 if len(row) >= 4 else 10.2)
            if i < len(widths):
                set_cell_width(cells[i], widths[i])
            if header and r_idx == 0:
                set_cell_shading(cells[i], "E8EEF5")


def add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    add_rich_text(p, caption, "微软雅黑", 10.5, bold=True, color="24415F")


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    add_rich_text(cap, caption, "宋体", 10, color="475569")


def _math_run(text: str):
    run = OxmlElement("m:r")
    text_el = OxmlElement("m:t")
    text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def _math_container(tag: str, children: list) -> object:
    container = OxmlElement(f"m:{tag}")
    for child in children:
        container.append(child)
    return container


def _math_sub(base: str, sub: str):
    el = OxmlElement("m:sSub")
    el.append(_math_container("e", [_math_run(base)]))
    el.append(_math_container("sub", [_math_run(sub)]))
    return el


def _math_sup(base, sup: str):
    el = OxmlElement("m:sSup")
    base_children = base if isinstance(base, list) else [base]
    el.append(_math_container("e", base_children))
    el.append(_math_container("sup", [_math_run(sup)]))
    return el


def _math_subsup(base: str, sub: str, sup: str):
    el = OxmlElement("m:sSubSup")
    el.append(_math_container("e", [_math_run(base)]))
    el.append(_math_container("sub", [_math_run(sub)]))
    el.append(_math_container("sup", [_math_run(sup)]))
    return el


def _math_rad(children: list):
    el = OxmlElement("m:rad")
    pr = OxmlElement("m:radPr")
    deg_hide = OxmlElement("m:degHide")
    deg_hide.set(qn("m:val"), "on")
    pr.append(deg_hide)
    el.append(pr)
    el.append(_math_container("e", children))
    return el


def _normalize_formula_text(text: str) -> str:
    return (
        text.replace("Omega_", "Ω_")
        .replace("theta_", "θ_")
        .replace("lambda", "λ")
        .replace("∑", "Σ")
        .replace("^2", "²")
        .replace("^T", "ᵀ")
    )


def _is_token_char(ch: str) -> bool:
    return (ch.isalnum() and ch not in "²ᵀ") or ch in "Ωηθλβμαγδ"


def _read_balanced(text: str, start: int) -> tuple[str, int]:
    depth = 0
    for pos in range(start, len(text)):
        if text[pos] == "(":
            depth += 1
        elif text[pos] == ")":
            depth -= 1
            if depth == 0:
                return text[start + 1 : pos], pos + 1
    return text[start + 1 :], len(text)


def _read_subscript(text: str, start: int) -> tuple[str, int]:
    delimiters = set(" +-*/=()[]{};<>≤≥∈∪·")
    pos = start
    while pos < len(text) and text[pos] not in delimiters:
        pos += 1
    return text[start:pos], pos


def _parse_math_text(text: str) -> list:
    text = _normalize_formula_text(text)
    elements: list = []
    i = 0
    while i < len(text):
        if text.startswith("sqrt(", i):
            inside, end = _read_balanced(text, i + 4)
            elements.append(_math_rad(_parse_math_text(inside)))
            i = end
            continue
        if text.startswith("log10", i):
            elements.append(_math_sub("log", "10"))
            i += 5
            continue
        ch = text[i]
        if ch == "²":
            if elements:
                prev = elements.pop()
                elements.append(_math_sup(prev, "2"))
            else:
                elements.append(_math_run(ch))
            i += 1
            continue
        if ch == "ᵀ":
            if elements:
                prev = elements.pop()
                elements.append(_math_sup(prev, "T"))
            else:
                elements.append(_math_run(ch))
            i += 1
            continue
        if _is_token_char(ch):
            j = i + 1
            while j < len(text) and _is_token_char(text[j]):
                j += 1
            base = text[i:j]
            if base == "Omega":
                base = "Ω"
            if base == "theta":
                base = "θ"
            sub = ""
            sup = ""
            if j < len(text) and text[j] == "_":
                sub, j = _read_subscript(text, j + 1)
            if j < len(text) and text[j] == "^":
                if j + 1 < len(text):
                    sup = text[j + 1]
                    j += 2
                else:
                    j += 1
            if sub and sup:
                elements.append(_math_subsup(base, sub, sup))
            elif sub:
                elements.append(_math_sub(base, sub))
            elif sup:
                elements.append(_math_sup(_math_run(base), sup))
            else:
                elements.append(_math_run(base))
            i = j
            continue
        elements.append(_math_run(ch))
        i += 1
    return elements


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    for element in _parse_math_text(text):
        math.append(element)
    math_para.append(math)
    p._p.append(math_para)


INLINE_MATH_PATTERNS = [
    r"R_EMI=exp\(-d/d0\)",
    r"R_EMI,stage",
    r"R_total,i",
    r"R_total",
    r"R_distance,i",
    r"R_height,i",
    r"R_Fresnel,i",
    r"R_Fresnel",
    r"R_weather",
    r"R_terrain,i",
    r"R_EMI",
    r"R_LoS,i",
    r"R_LoS",
    r"R_pre,i",
    r"R_lim",
    r"S_weather",
    r"S_equipment",
    r"S_humidity",
    r"S_tower",
    r"S_hover",
    r"S_dist",
    r"S_req",
    r"R_rain",
    r"H_humidity",
    r"V_fog",
    r"i_F\(t\)",
    r"i_B\(t\)",
    r"y\(t\)",
    r"s\(t\)",
    r"n\(t\)",
    r"h_LoS\(x\)",
    r"h_DEM\(x\)",
    r"h_DSM\(x\)",
    r"h_UAV",
    r"h_min",
    r"h_max",
    r"h_G",
    r"h_i",
    r"d_tower,i",
    r"d_bs,i",
    r"d_safe",
    r"d_min",
    r"d_3D",
    r"d_h",
    r"d0",
    r"q_coord,i",
    r"q_height,i",
    r"q_DEM",
    r"q_link",
    r"q_i",
    r"z_i",
    r"η_i",
    r"Ω",
    r"Omega_sandbox",
    r"Relation_i",
    r"theta_assumption",
    r"L_FSPL",
    r"L_extra",
    r"P_t",
    r"G_t",
    r"G_r",
    r"C\(x\)>0",
    r"C\(x\)≤0",
    r"C\(x\)",
    r"C_min",
    r"C_F",
    r"C0",
    r"B_LoS=0",
    r"B_LoS=1",
    r"B_LoS",
    r"0\.6F1",
    r"F1",
    r"λ",
    r"β_r",
    r"β_h",
    r"β_v",
    r"μ1",
    r"μ2",
    r"μ3",
    r"μ4",
    r"α",
    r"β",
    r"γ",
    r"δ",
    r"E_flight",
    r"T_task",
    r"Q_image",
    r"J",
    r"v · r_i\^T",
    r"r_i",
]

INLINE_MATH_RE = re.compile("|".join(f"(?:{pattern})" for pattern in sorted(INLINE_MATH_PATTERNS, key=len, reverse=True)))


def _append_inline_math(paragraph, text: str) -> None:
    math = OxmlElement("m:oMath")
    for element in _parse_math_text(text):
        math.append(element)
    paragraph._p.append(math)


def add_rich_text(
    paragraph,
    text: str,
    font_name: str = "宋体",
    size: float = 12,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    pos = 0
    for match in INLINE_MATH_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_east_asia_font(run, font_name)
            run.font.size = Pt(size)
            if bold is not None:
                run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        _append_inline_math(paragraph, match.group(0))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_east_asia_font(run, font_name)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_body_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_rich_text(p, text, "宋体", 12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_rich_text(p, text, "宋体", 11.5)


def build_docx(summary: dict[str, object], figures: dict[str, Path]) -> None:
    doc = Document()
    style_doc(doc)
    add_title_page(doc, summary)

    # Section 1.
    doc.add_heading("1. 研究任务、数据基础与技术范围", level=1)
    add_body_paragraph(
        doc,
        "本报告面向泰山及泰安区域低空无人机巡检场景，重点分析飞行高度、电力设备电磁干扰、山区地形遮挡与通信稳定性之间的影响机制。"
        "报告所用项目数据来自已标准化处理的航线、航点、杆塔、基站候选点和三维仿真沙盘展示结果；传播模型来自公开推荐书、论文和工程常用计算公式。"
        "研究目标在于构建由“数据清点—传播机理—风险识别—策略优化”组成的技术论证链条，为后续实地测试、设备选型和任务参数整定提供可复用的分析框架。"
    )
    add_body_paragraph(
        doc,
        "已有数据可验证的部分包括：三条 KML 航线共 72 个航点，航点具有经纬度和 KML 高度字段；线路台账中有 849 条杆塔记录，其中 831 条具备有效经纬度；"
        "任务统计表中有 1414 条任务记录和 10 个基站候选点。KML 高度字段在不同航线文件中呈现明显差异，本文将其作为“航点高度/航线设定高度候选”纳入传播路径分析，"
        "并把高度含义、地形基准和设备天线参数列为后续校准条件。"
    )
    add_body_paragraph(
        doc,
        "从研究范式看，低空巡检通信链路具有显著的空间耦合特征：航线高度决定几何视距与斜距，山区地形决定遮挡和绕射条件，杆塔及导线结构改变近场电磁暴露与多径环境，"
        "气象因素进一步调制传播衰减和电晕/局放活动。上述因素共同作用，使通信可靠性呈现出“高度—距离—地形—电磁环境”多变量耦合特征。"
    )
    add_body_paragraph(
        doc,
        "本文采用分层研究逻辑：首先以现有数据建立可追溯的空间对象体系；其次以 FSPL、LoS/NLoS、第一菲涅尔区和 EMI 暴露模型解释关键影响机制；"
        "最后将各因素归一化为综合风险分量，形成可用于任务规划、航线复核和测试方案设计的优化策略。"
    )

    add_table_caption(doc, "表 1 合同要求—本章研究内容—当前实现方式")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("合同要求", "本报告对应内容", "当前实现方式"),
            ("分析飞行高度对传播路径、衰减、覆盖和通信质量的影响", "第 2 节建立 LoS/NLoS、斜距、FSPL 与第一菲涅尔区净空机理。", "采用理论公式和模型示意，给出 920 MHz、2.4 GHz、5.8 GHz 在 3 km、4 km 的 FSPL 表。"),
            ("分析电力设备电磁干扰对通信链路的影响", "第 4 节拆分电晕、局放、工频场、暂态、金属反射散射及耦合路径。", "建立 R_EMI 暴露风险表达式，并给出现场频谱监测与抗扰度校准条件。"),
            ("提出通信稳定性和可靠性优化策略", "第 5 节建立 R_total 综合风险模型并给出低/中/高/极高分级。", "围绕高度、节点/中继、链路冗余、近塔巡检和失联保护给出策略表。"),
            ("结合 DEM、三维沙盘和航线杆塔数据", "第 3 节给出 DEM 剖面、LoS 判定、0.6F1 净空和 3–4 km 敏感区流程。", "三维沙盘用于空间叠加和航线复核，DEM/DSM 数据用于后续地形剖面计算。"),
        ],
        [Cm(4.0), Cm(6.2), Cm(6.2)],
        header=True,
    )

    add_table_caption(doc, "表 2 数据基础、可用信息与后续校准条件")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    route_ranges = "；".join(
        f"{item['name']}：{item['waypoint_count']} 点，{item['min_height']:.1f}–{item['max_height']:.1f} m"
        for item in summary["route_stats"]
    )
    fill_table(
        table,
        [
            ("数据/资源", "当前数量或字段", "可用信息", "后续校准条件"),
            ("KML 航线与航点", route_ranges, "航点经纬度、KML 高度、航点数、航线长度候选。", "需结合航线任务说明确认轨迹完整性、起飞基准和高度含义。"),
            ("线路台账杆塔", f"{len(summary['towers'])} 条记录，{len(summary['valid_towers'])} 条有效经纬度。", "杆塔经纬度、杆塔高度、线路名、杆号、与航线最近距离。", "需补充线路拓扑、杆塔连接顺序及缺坐标杆塔的人工复核结果。"),
            ("基站候选点", f"{len(summary['stations'])} 个，含经纬度和海拔字段。", "可计算航点到基站候选点的几何距离。", "需补充天线挂高、频段、发射功率、增益、下倾角和接收灵敏度。"),
            ("三维仿真沙盘", "React + CesiumJS 原型，叠加航线、杆塔、基站和航点。", "可验证图层叠加、对象属性和当前视觉状态。", "用于空间关系复核；地形高程计算采用独立 DEM/DSM 数据接口。"),
            ("DEM/DSM", "目录与接口预留；当前处于待接入状态。", "可形成地形剖面采样、LoS 判定和净空计算流程。", "需明确数据源、分辨率、坐标系、高程基准和插值方法。"),
            ("通信质量数据", "RSSI、SNR、丢包率、图传码率和断链位置等观测量待采集。", "可先建立理论链路预算和风险分级框架。", "后续通过飞行日志、频谱监测和链路测试完成模型校准。"),
        ],
        [Cm(3.0), Cm(4.3), Cm(4.4), Cm(4.7)],
        header=True,
    )
    add_body_paragraph(
        doc,
        "网页端仿真原型已形成多地图模式展示能力。地形影像模式可在较大尺度上检查三条 KML 航线簇与杆塔空间分布，适合用于巡检线路覆盖范围、航点外包络和台账对象位置的一致性复核；"
        "拓扑底图模式更强调道路、地名和地物参照关系，便于运维人员将线路空间位置与实际巡检组织方式对应；三维模式则用于局部航段的高度、坡面和杆塔相对关系观察。"
        "上述截图共同构成报告的空间证据链，使传播模型分析与项目已有航线、杆塔和可视化成果保持一致。"
    )
    add_figure(
        doc,
        figures["web_terrain"],
        "图 1 网页端地形影像模式总览。图中叠加三条 KML 航线簇、航点和杆塔空间对象，可用于检查巡检线路覆盖范围与台账空间分布的一致性。",
        width_in=6.55,
    )
    add_figure(
        doc,
        figures["web_comparison"],
        "图 2 网页端多地图模式对照。地形影像、拓扑底图、三维显示和基础网格降级模式分别服务于线路覆盖核查、地物参照、局部空间关系复核和底图不可用时的系统状态说明。",
        width_in=6.55,
    )
    add_body_paragraph(
        doc,
        "多地图模式并不是简单的视觉切换，而是面向不同分析任务的空间表达方式。对于通信链路研究，地形影像模式优先用于判别航线是否处在山脊、谷地、坡面或开阔区域；"
        "拓扑底图模式用于辅助识别道路可达性、地面站布设候选位置和线路走廊；三维模式用于确认局部航段与山体起伏、杆塔高度和基站候选点的空间关系。"
        "基础网格模式保留了外部地图服务不可用时的图层降级能力，有助于区分数据渲染问题与底图加载问题。"
    )

    # Section 2.
    doc.add_page_break()
    doc.add_heading("2. 无人机飞行高度对通信传播路径、衰减特性与覆盖范围的影响机制", level=1)
    add_body_paragraph(
        doc,
        "飞行高度首先影响视距链路（Line-of-Sight, LoS）是否成立。山区低空巡检中，山体、山脊、线路走廊起伏和杆塔/导线金属结构均可能把 LoS 转为非视距链路（Non-Line-of-Sight, NLoS），"
        "此时接收信号会叠加绕射、反射、散射和多径衰落。高度上升通常有助于越过局部遮挡并扩大几何覆盖范围，但也会增加无人机与地面站或基站之间的斜距。"
        "因此，高度不是越高越好，而是在“改善视距/净空”和“增加路径损耗/飞行安全约束”之间折中。"
    )
    add_body_paragraph(
        doc,
        "从几何关系看，若地面站天线高度和水平距离保持不变，飞行高度增加会抬升传播直线，降低山脊、林带和塔材遮挡概率；"
        "同时，链路斜距随高度增加而增大，FSPL 也随距离对数增长。对低空巡检而言，最优高度往往并非单一固定值，而是沿航线地形、塔位和任务载荷需求变化的分段高度策略。"
    )
    add_body_paragraph(
        doc,
        "第一菲涅尔区净空是高度优化中的关键约束。即使几何 LoS 未被完全遮挡，若 0.6F1 区域被山体或金属结构侵入，链路仍可能出现绕射损耗和多径衰落。"
        "因此，工程上应同时考察“LoS 是否成立”和“0.6F1 是否具备足够净空”，再把图传、遥控和数传链路的频段差异纳入链路预算。"
    )
    add_figure(
        doc,
        figures["height"],
        "图 3 不同飞行高度下传播路径、山体遮挡和第一菲涅尔区净空示意。低高度易产生 NLoS；中等高度改善净空；过高高度虽然视距更好，但斜距和路径损耗增加。",
        width_in=6.25,
    )
    add_body_paragraph(
        doc,
        "自由空间路径损耗（Free Space Path Loss, FSPL）用于描述理想自由空间中电波随频率和距离增加而产生的理论损耗。"
        "该指标适合作为链路预算的基础项，用于比较不同频段和距离条件下的理论损耗差异。工程计算式为："
    )
    add_formula(doc, "L_FSPL = 92.45 + 20 log10(f_GHz) + 20 log10(d_km)")
    add_body_paragraph(
        doc,
        "其中 f_GHz 为频率（GHz），d_km 为链路距离（km）。表 3 给出 920 MHz、2.4 GHz 和 5.8 GHz 在 3 km 与 4 km 下的理论 FSPL。"
        "可以看到，频率越高、距离越远，自由空间损耗越大；在同一距离下，5.8 GHz 相比 920 MHz 的理论损耗约高 16 dB。"
    )
    add_table_caption(doc, "表 3 3 km 和 4 km 下不同频段理论 FSPL")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fspl_rows = [("频段", "3 km FSPL (dB)", "4 km FSPL (dB)", "工程解释")]
    for freq, label in [(0.92, "920 MHz"), (2.4, "2.4 GHz"), (5.8, "5.8 GHz")]:
        fspl_rows.append((label, f"{fspl_db(freq, 3):.2f}", f"{fspl_db(freq, 4):.2f}", "用于链路预算基础项；实际余量需叠加天线、地形、多径、干扰和接收机参数。"))
    fill_table(table, fspl_rows, [Cm(3.1), Cm(3.1), Cm(3.1), Cm(7.1)], header=True)
    add_body_paragraph(
        doc,
        "如图 4 所示，MATLAB 按上述公式生成了频段—距离损耗曲线。曲线表明，在 0.5–5 km 范围内，路径损耗随距离呈对数增长；"
        "同一距离条件下，高频链路相对低频链路具有更高的自由空间损耗。3–4 km 区间被标注为工程敏感距离区间，是因为该范围内距离损耗、地形遮挡、菲涅尔区净空和近塔 EMI 暴露更容易发生叠加，"
        "适合作为任务前链路预算与航线复核的重点分析窗口。"
    )
    add_figure(
        doc,
        figures["matlab_fspl"],
        "图 4 MATLAB 理论计算得到的 920 MHz、2.4 GHz 和 5.8 GHz 自由空间路径损耗曲线，距离范围 0.5–5 km，并标出 3–4 km 工程敏感距离区间。",
        width_in=6.25,
    )
    add_body_paragraph(
        doc,
        "菲涅尔区曲线进一步揭示频段选择与净空要求之间的矛盾关系。较低频段具有更低 FSPL，但波长更长，第一菲涅尔区半径也更大，对山体、树冠和塔材结构的净空要求更高；"
        "较高频段菲涅尔区半径相对较小，但路径损耗和多径敏感性增强。因而在山区巡检链路设计中，频段选择不能只看理论覆盖距离，还应同步评估地形剖面、航点高度和天线安装条件。"
    )
    add_figure(
        doc,
        figures["matlab_fresnel"],
        "图 5 MATLAB 模型计算得到的第一菲涅尔区 F1 与 0.6F1 随链路距离变化曲线。该图用于说明不同频段在净空判据上的差异。",
        width_in=6.25,
    )
    add_body_paragraph(
        doc,
        "覆盖范围应理解为链路预算、天线方向图、地形遮挡和业务质量门限共同作用的结果。FSPL 给出距离与频段的基础损耗，天线增益和接收灵敏度决定可用链路余量，"
        "地形和塔材结构决定 LoS/NLoS 转换，多径与干扰则改变短时可靠性。本文建议将高度相关风险拆分为 R_height、R_distance、R_LoS 和 R_Fresnel 等可解释分量，"
        "以便后续用 RSSI、SNR、丢包率和图传码率等观测量进行参数标定。"
    )
    add_body_paragraph(
        doc,
        "对泰山山区巡检场景，飞行高度优化还需要兼顾图像采集精度、杆塔安全距离、空域限制、返航能耗和任务持续时间。"
        "当航点位于山谷或坡面背向区域时，适度提高高度可能显著改善 LoS；当航线已处于开阔视距区域时，继续升高带来的净空收益会逐渐降低，而链路损耗、能耗和安全控制成本会相应增加。"
    )

    # Section 3.
    doc.add_page_break()
    doc.add_heading("3. 基于 DEM 与三维仿真沙盘的山区通信敏感区分析", level=1)
    add_body_paragraph(
        doc,
        "当前三维 demo 已能把航线、航点、杆塔和基站候选点叠加在三维场景中，用于核查空间关系和对象属性。"
        "后续一旦接入真实数字高程模型（Digital Elevation Model, DEM）或数字表面模型（Digital Surface Model, DSM），即可沿地面站/基站到无人机航点的剖面采样，"
        "判断 LoS 是否被地形截断，并计算第一菲涅尔区（First Fresnel Zone）净空。"
    )
    add_body_paragraph(
        doc,
        "山区通信敏感区的识别应以“剖面”为基本分析单元，而不是只依赖二维平面距离。对任意航点，可从候选地面站或通信节点出发，沿大地线或局部投影直线提取 DEM 高程序列；"
        "再将地形剖面与无人机天线高度、地面站天线高度和频段波长进行组合，得到 LoS 直线、F1 半径和净空余量。"
    )
    add_body_paragraph(
        doc,
        "三维仿真沙盘在该流程中承担空间索引和可视化复核作用：一方面，它可以展示航线、杆塔、航点和基站候选点的相对位置，便于确定需要抽取剖面的关键航段；"
        "另一方面，它可以把风险分级结果叠加到空间对象上，形成面向运维人员的航线复核界面。DEM/DSM 则承担定量高程输入职责，两者在工程系统中应保持数据职责分离。"
    )
    add_body_paragraph(
        doc,
        "网页端较大尺度线路覆盖视图与局部三维空间视图分别见图 6 和图 7。前者强调三条 KML 航线簇在地图底座上的整体分布，有助于识别线路覆盖范围、航点聚集区和杆塔空间邻近关系；"
        "后者强调局部航点、杆塔和三维地形之间的相对位置，可作为选择 DEM 剖面方向、识别遮挡候选点和解释局部传播路径变化的辅助依据。"
        "在工程流程中，二维总览用于筛选重点航段，三维视图用于复核局部空间机制，二者与 DEM 剖面计算共同构成山区通信敏感区分析链条。"
    )
    add_figure(
        doc,
        figures["web_topographic"],
        "图 6 网页端拓扑/地形底图模式下的巡检线路覆盖总览。该视图用于复核三条 KML 航线簇、航点和杆塔对象的整体空间关系。",
        width_in=6.55,
    )
    add_figure(
        doc,
        figures["web_3d"],
        "图 7 网页端三维模式下的局部航线、杆塔和地形关系。该视图适合用于解释局部高度变化、遮挡候选位置和近塔空间关系。",
        width_in=6.55,
    )
    add_body_paragraph(
        doc,
        "建议流程为：航点/地面站坐标 → DEM 剖面采样 → LoS 判定 → 第一菲涅尔区净空 → 3–4 km 敏感区识别 → 通信风险分级。"
        "其中，3–4 km 定义为工程敏感距离区间和仿真分析区间，后续可结合设备链路预算和现场测试数据进行阈值校准。"
    )
    add_formula(doc, "h_LoS(x) = h_G + x/D · (h_UAV - h_G)")
    add_formula(doc, "C(x) = h_LoS(x) - h_DEM(x) - 0.6F1(x)")
    add_body_paragraph(
        doc,
        "式中 h_G 为地面站或基站天线高度，h_UAV 为无人机天线高度，D 为总链路距离，x 为剖面距离。"
        "C(x) 为考虑 0.6 倍第一菲涅尔区后的净空余量；当 C(x)>0 时可认为该剖面点满足工程净空，当 C(x)≤0 时说明 0.6F1 被地形侵入或 LoS 可能断裂。"
    )
    add_body_paragraph(
        doc,
        "该判据具有明确的工程可解释性：h_LoS(x) 描述传播直线，h_DEM(x) 描述地形剖面，0.6F1(x) 描述为降低绕射损耗而预留的净空要求。"
        "当 C(x) 的最小值接近零时，即使链路在几何上仍处于视距状态，也应纳入中风险复核；当 h_DEM(x) 高于 LoS 直线时，则需要通过抬高航点、调整通信节点或设置中继来恢复传播条件。"
    )
    add_figure(
        doc,
        figures["dem"],
        "图 8 DEM 地形剖面与 0.6F1 净空示意。该图说明 LoS 直线、地形剖面和 0.6F1 净空边界之间的几何判据关系。",
        width_in=6.25,
    )
    add_body_paragraph(
        doc,
        "MATLAB 剖面示意（图 9）将地形、LoS 直线和 0.6F1 净空边界放在同一坐标系中，突出 C(x) 的空间含义。"
        "当地形曲线接近或穿越 0.6F1 净空边界时，链路即使未出现完全遮挡，也会进入绕射和多径更加敏感的区段。"
        "该类剖面图适合用于解释“为什么某一航段需要调高航点、改变地面站位置或增加中继节点”，从而把抽象风险指数转化为可复核的空间原因。"
    )
    add_figure(
        doc,
        figures["matlab_dem"],
        "图 9 MATLAB DEM 地形剖面与 0.6F1 净空判据示意。地形、LoS 和净空边界共同用于识别遮挡与净空不足区段。",
        width_in=6.25,
    )
    add_body_paragraph(
        doc,
        "飞行高度变化对最小净空余量的影响见图 10。该曲线的学术意义在于把高度优化从经验判断转化为剖面约束问题："
        "当最小 C(x) 随高度增加由负值转为正值时，说明高度调整正在消除净空不足；当曲线进入平台区后，继续升高对净空的边际改善降低，而链路斜距、能耗和空域约束的代价上升。"
        "因此，合理高度应位于净空满足、链路损耗可控和巡检成像质量可接受的综合平衡区间。"
    )
    add_figure(
        doc,
        figures["matlab_height_margin"],
        "图 10 MATLAB 模型得到的飞行高度与最小净空余量关系。曲线用于识别满足 0.6F1 净空的高度阈值和边际收益变化。",
        width_in=6.25,
    )
    add_body_paragraph(
        doc,
        "距离—高度二维热力图（图 11）表达了综合通信风险的空间变化趋势。热力图把距离损耗、视距遮挡和净空不足等因素投影到同一分析平面，"
        "可用于查找高风险组合区、比较不同高度策略的敏感性，并为航线重规划提供直观依据。"
        "在后续工程化应用中，该图可由实际 DEM 剖面、设备链路预算和日志观测量重新计算，从而形成随数据更新的风险图层。"
    )
    add_figure(
        doc,
        figures["matlab_risk_heatmap"],
        "图 11 MATLAB 距离—高度二维通信风险热力图。颜色变化表示综合风险指数的相对大小，用于分析距离与高度耦合下的敏感区。",
        width_in=6.25,
    )
    add_table_caption(doc, "表 4 LoS 与 0.6F1 净空风险分级")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("风险等级", "判据", "工程含义", "校准与应用条件"),
            ("低风险", "LoS 清晰，且 C(x)>0，全剖面满足 0.6F1。", "传播路径具备较好的几何净空，链路余量主要受频段和设备参数控制。", "接入 DEM/DSM 后可作为常规航段基准。"),
            ("中风险", "LoS 未断，但部分剖面 C(x)≤0。", "菲涅尔区受到侵入，可能产生附加绕射损耗和链路余量下降。", "需复核高度、频段和天线姿态。"),
            ("高风险", "山体或障碍物遮挡 LoS。", "NLoS 条件显著增强，多径、绕射和链路不稳定性上升。", "优先考虑抬升航点、调整节点或设置中继。"),
            ("通信敏感区", "高风险地形叠加 3–4 km 工程敏感距离区间。", "距离损耗和遮挡风险同时上升，是任务前链路验证重点。", "需 DEM、链路预算和飞行日志联合校准。"),
        ],
        [Cm(2.3), Cm(5.0), Cm(4.8), Cm(4.3)],
        header=True,
    )
    add_body_paragraph(
        doc,
        "在数据管理上，在线底图、影像底图和三维显示底座宜限定为可视化背景；地形剖面计算应使用可追溯的 DEM/DSM 数据源。"
        "后续若导入 SRTM、地方测绘 DEM 或无人机倾斜摄影 DSM，应记录数据来源、分辨率、坐标系、高程基准和插值方法，并将剖面采样结果写入 `data/processed/`，以保证分析过程可复现。"
    )
    add_body_paragraph(
        doc,
        "对于当前三条航线，航点数量较少且空间半径较小，更适合作为单塔或局部巡检航线样本。"
        "因此，敏感区分析的工程价值主要体现在方法验证和接口预留：先证明航点、杆塔、基站候选点能够形成统一空间索引，再在 DEM 接入后扩展到跨塔、跨山脊和长距离巡检航段。"
    )

    # Section 4.
    doc.add_page_break()
    doc.add_heading("4. 电力设备电磁干扰及气象耦合因素对通信链路可靠性的影响", level=1)
    add_body_paragraph(
        doc,
        "电力巡检场景中，电磁干扰（Electromagnetic Interference, EMI）主要来自高压导线电晕放电、绝缘子和金具局部放电、工频电场/磁场、变电或换流设备开关瞬态，以及杆塔和导线金属结构引起的反射与散射。"
        "这些因素会对遥控链路、图传链路、数传链路、GNSS/北斗、磁罗盘、飞控系统和任务载荷产生不同形式的影响。"
    )
    add_body_paragraph(
        doc,
        "电力设备 EMI 对无人机链路的影响具有频域、时域和空间域三重特征。频域上，电晕和局放可能抬升宽带噪声底或形成脉冲型干扰；"
        "时域上，开关暂态和放电活动可能引起短时误码、链路重传或图传瞬断；空间域上，杆塔、导线和金具构成复杂金属散射体，会改变无人机近塔悬停时的多径结构和天线方向图有效性。"
    )
    add_body_paragraph(
        doc,
        "在巡检任务中，遥控链路通常承担飞行控制安全底线，图传链路承担视觉巡检质量，数传链路承担任务状态回传，GNSS/北斗和磁罗盘则影响定位与姿态估计。"
        "因此，EMI 分析不应只聚焦单一通信通道，而应从“通信链路—导航传感器—飞控系统—任务载荷”的系统链条评估风险传导路径。"
    )
    add_body_paragraph(
        doc,
        "从耦合路径看，前门耦合（Front-door Coupling）是指干扰经通信或导航天线进入接收机前端，表现为同频/邻频干扰、噪声底抬升、灵敏度下降或误码率上升；"
        "后门耦合（Back-door Coupling）是指干扰经机体结构、线缆、接口、电源或传感器通道进入系统，可能引发磁罗盘异常、GNSS 质量下降、飞控姿态估计扰动或载荷数据异常。"
    )
    add_figure(
        doc,
        figures["emi"],
        "图 12 电力设备电磁干扰对无人机通信链路影响路径图。R_EMI 用于描述电磁暴露风险排序，后续可通过频谱监测、设备抗扰度试验和飞行日志进行参数校准。",
        width_in=6.25,
    )
    add_formula(doc, "R_EMI = a1S_dist + a2S_tower + a3S_weather + a4S_equipment")
    add_body_paragraph(
        doc,
        "式中 S_dist 可表示与杆塔、导线或设备的近距离暴露，S_tower 表示杆塔/导线结构反射散射风险，S_weather 表示雨、雾、湿度等气象修正项，S_equipment 表示设备类型和开关状态等因素。"
        "雨、雾和高湿度既可能引入额外传播衰减，也可能增强电晕或局部放电活动，因此可作为气象耦合修正项纳入后续参数标定。"
    )
    add_body_paragraph(
        doc,
        "R_EMI 的工程价值在于把难以直接统一量纲的电磁因素转化为可排序、可叠加、可校准的暴露指标。"
        "在近塔巡检阶段，S_dist 可按无人机与杆塔/导线的最近距离构造衰减函数；S_tower 可由塔材密度、悬停姿态和天线遮挡关系近似；"
        "S_weather 可由降雨、湿度和能见度修正；S_equipment 则可反映变电、换流或开关设备的运行状态差异。"
    )
    add_body_paragraph(
        doc,
        "近塔距离衰减曲线（图 13）采用 R_EMI=exp(-d/d0) 形式展示近塔距离与电磁暴露风险之间的相对关系，其中 d 为无人机到杆塔、导线或设备的最近距离，d0 表示场景尺度参数。"
        "d0 取 50 m、80 m 和 120 m 时，曲线衰减速度不同，分别对应局部源主导、线路走廊主导和较宽空间暴露主导等不同建模假设。"
        "该图的意义在于说明近塔距离控制对通信和导航系统的风险削减作用，并为设置最小安全距离、悬停时间和任务姿态约束提供量化表达入口。"
    )
    add_figure(
        doc,
        figures["matlab_emi_distance"],
        "图 13 MATLAB 工程假设下的电磁干扰暴露风险随近塔距离变化曲线，R_EMI=exp(-d/d0)，d0 分别取 50 m、80 m 和 120 m。",
        width_in=6.25,
    )
    add_table_caption(doc, "表 5 EMI 风险源、影响对象与校准条件")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("风险源", "主要影响对象", "可能表现", "校准条件"),
            ("导线电晕/局放", "遥控、图传、数传、GNSS", "噪声底抬升、误码率上升、定位质量下降。", "需频谱监测、天气记录和缺陷状态记录。"),
            ("工频电场/磁场", "磁罗盘、飞控、载荷线缆", "姿态估计扰动、传感器读数异常。", "需现场电磁环境和机型抗扰度数据。"),
            ("开关暂态", "接收机前端、电源与接口", "瞬态误码、链路重传、设备保护动作。", "需设备运行状态记录。"),
            ("金属结构反射/散射", "图传/数传链路", "多径衰落、短时吞吐下降。", "需频段、天线姿态和环境建模。"),
            ("雨雾湿度", "传播路径与放电活动", "衰减修正或电晕活动增强。", "需气象观测与链路日志同步采样。"),
        ],
        [Cm(3.1), Cm(3.6), Cm(4.5), Cm(5.2)],
        header=True,
    )

    # Section 5.
    doc.add_page_break()
    doc.add_heading("5. 通信链路优化策略与无人机巡检安全控制建议", level=1)
    add_body_paragraph(
        doc,
        "综合风险模型应把高度、距离、视距、菲涅尔区、电磁暴露和天气因素分开建模，再按任务阶段和设备链路预算设置权重。"
        "本报告建议的工程预评估表达式如下，其中各分量均为归一化风险指数，权重 wi 可通过设备参数、专家规则、飞行日志和链路观测量逐步校准。"
    )
    add_body_paragraph(
        doc,
        "该模型的核心思想是把复杂环境中的通信可靠性问题转化为可解释的分量叠加。R_height 表征航点高度相对地形和任务需求的适配性，R_distance 表征链路斜距和 FSPL 压力，"
        "R_LoS 与 R_Fresnel 表征地形遮挡与净空条件，R_EMI 表征近塔电磁暴露，R_weather 表征降雨、湿度、风速等环境修正。"
        "通过分量化建模，运维人员可以识别风险来源，而不仅是获得一个综合等级。"
    )
    add_formula(doc, "R_total = w1R_height + w2R_distance + w3R_LoS + w4R_Fresnel + w5R_EMI + w6R_weather")
    add_figure(
        doc,
        figures["risk"],
        "图 14 通信链路综合风险分级与优化策略流程图。该流程面向任务预评估、航线复核和后续系统接口设计，可与现场安全评估流程协同使用。",
        width_in=6.25,
    )
    add_table_caption(doc, "表 6 综合通信风险分级建议")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("等级", "R_total 区间", "典型条件", "控制建议"),
            ("低风险", "R < 0.25", "LoS 清晰、0.6F1 净空满足、距离较短、EMI 暴露低。", "按常规航线执行，保留链路日志。"),
            ("中风险", "0.25–0.50", "距离偏大或净空裕度不足，天气/近塔暴露轻度叠加。", "复核高度、速度、返航点和备用链路。"),
            ("高风险", "0.50–0.75", "LoS 遮挡、3–4 km 敏感区叠加近塔或气象因子。", "调整航高或通信节点，必要时设置中继。"),
            ("极高风险", "R ≥ 0.75", "遮挡、长距离、EMI、恶劣天气多因素耦合。", "重规划任务参数，优先开展链路验证和应急策略演练。"),
        ],
        [Cm(2.3), Cm(2.6), Cm(6.5), Cm(5.0)],
        header=True,
    )
    add_table_caption(doc, "表 7 风险问题—优化策略—预期效果")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("风险问题", "优化策略", "预期效果"),
            ("低高度受山体、林带或线路走廊起伏遮挡", "在满足空域和巡检安全的前提下提高局部航点高度，并用 DEM 剖面复核 LoS 和 0.6F1。", "提升视距概率和净空余量，降低 NLoS 风险。"),
            ("高度过高导致斜距和路径损耗增大", "采用分段高度策略，避免全程固定高高度；关键航段以链路预算和影像任务需求共同约束。", "在覆盖范围与链路损耗之间取得平衡。"),
            ("3–4 km 工程敏感区距离损耗上升", "优化地面站/基站位置，必要时设置临时中继或车载中继；任务前验证返航与断链保护。", "缩短有效链路距离，增加链路预算余量。"),
            ("近塔 EMI 暴露与金属多径", "限定近塔悬停时间和最小安全距离，减少天线被塔材遮挡；采用抗干扰编码、跳频或双链路备份。", "降低短时链路波动和控制链路失效概率。"),
            ("GNSS/北斗和磁罗盘易受干扰", "靠近高压设备时提高视觉定位、惯导和航迹约束权重；设置异常姿态和定位质量触发规则。", "降低定位异常向飞控风险传导。"),
            ("通信质量闭环待完善", "任务日志中补采 RSSI、SNR、丢包率、图传码率、频谱监测和断链位置，并与航点、天气、设备状态对齐。", "将预评估模型逐步校准为可验证的链路可靠性模型。"),
        ],
        [Cm(4.2), Cm(7.0), Cm(5.2)],
        header=True,
    )
    add_body_paragraph(
        doc,
        "综合来看，通信链路优化应遵循“先识别主导风险、再选择最小代价控制措施”的原则。"
        "若主导风险来自地形遮挡，应优先调整高度或通信节点；若主导风险来自长距离损耗，应优先优化地面站位置或设置中继；"
        "若主导风险来自近塔 EMI，应优先控制悬停距离、姿态和链路冗余；若多因素耦合，则应重规划航线并开展专项链路验证。"
    )
    add_body_paragraph(
        doc,
        "对后续工程实施，建议建立三类闭环：第一，数据闭环，将 KML、杆塔、基站、DEM 和通信日志统一到航点尺度；第二，模型闭环，将 FSPL、LoS、Fresnel 和 EMI 指标与观测量对齐；"
        "第三，控制闭环，把风险等级映射到飞行高度、速度、通信节点、中继策略和失联保护参数。"
        "通过上述闭环，通信可靠性分析可以从静态报告逐步演进为可更新、可验证、可服务任务调度的工程工具。"
    )
    doc.add_heading("参考依据摘要与复核提示", level=2)
    add_body_paragraph(
        doc,
        "本报告参考 ITU-R P.525/P.526/P.530、UAV A2G 信道综述、低空平台覆盖高度模型、SRTM/NASA Earthdata 资料以及高压线路电磁干扰相关标准和文献。"
        "完整来源、模型假设和校准条件见同目录 `report_notes.md`。对论文卷期页码、标准版本和 DOI 的最终引用，建议在正式对外交付前由人工再次核验。"
    )

    doc.save(REPORT_PATH)


def add_reference_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    add_rich_text(p, text, "宋体", 10.5)


def build_docx_revised(summary: dict[str, object], figures: dict[str, Path]) -> None:
    doc = Document()
    style_doc(doc)
    add_title_page(doc, summary)

    # Section 1.
    doc.add_heading("1. 研究任务、数据基础与技术边界", level=1)
    add_body_paragraph(
        doc,
        "本报告围绕泰山低空巡检无人机通信链路可靠性开展机理分析与方案预评估，研究对象包括飞行高度、地形遮挡、第一菲涅尔区（First Fresnel Zone）净空、"
        "自由空间路径损耗（Free Space Path Loss, FSPL）以及电力设备电磁干扰（Electromagnetic Interference, EMI）等关键因素。"
        "报告以现有航线、航点高度、杆塔空间数据、基站候选点、KML 文件和三维仿真沙盘为基础，形成面向后续测试与工程优化的分析框架。"
    )
    add_body_paragraph(
        doc,
        "本报告重点覆盖合同研究内容 3 中关于无人机飞行高度、电力设备电磁干扰与通信链路可靠性的部分；航程/能耗与多气象因素的关联仅作为边界条件和后续扩展方向，不在本报告中展开实测分析。"
        "当前无 RSSI、SNR、丢包率、图传码率、频谱监测和真实断链位置等观测量，因此本文将传播模型、空间邻近性和工程假设作为预评估依据，并明确列出后续校准条件。"
    )
    add_body_paragraph(
        doc,
        "当前阶段基于 DEM/DSM 方法框架和示意地形剖面进行预评估，真实 DEM 接入后可复算 LoS、0.6F1 净空和敏感区。"
        "在线地图、影像底图和三维显示仅作为可视化平台底座，地形计算应使用可追溯的数字高程模型（Digital Elevation Model, DEM）或数字表面模型（Digital Surface Model, DSM）。"
    )
    add_body_paragraph(
        doc,
        "从研究方法上看，三维仿真沙盘并不只是报告插图来源，而是把航线、航点、杆塔、基站候选点与地形背景组织到同一空间语义中的分析载体。"
        "其价值在于把通信链路问题从抽象的距离和频段计算，转化为具有地理位置、相对高度、线路走廊和近塔关系的空间问题。"
        "因此，本文后续关于高度、遮挡、净空和 EMI 暴露的讨论，均以沙盘提供的空间叠加关系作为场景约束，以传播模型作为机理解释，以后续实测数据作为校准入口。"
    )
    add_body_paragraph(
        doc,
        "在可视化沙盘中，航点与杆塔对象的叠加能够揭示两个层面的信息：一是航线是否围绕局部塔位形成紧凑巡检轨迹，二是航点与基站候选点、山体起伏和线路走廊之间是否存在潜在传播敏感关系。"
        "这种空间判读不能替代通信质量观测，但可为后续测试布点、DEM 剖面抽取、近塔安全距离设置和链路冗余策略选择提供先验依据。"
        "换言之，沙盘模拟给出的不是确定性链路性能，而是使各类风险因子能够被定位、排序和复核的空间框架。"
    )
    add_body_paragraph(
        doc,
        "为了使上述空间框架具有可计算性，本文将沙盘中的对象抽象为多源空间对象集合 Ω。"
        "其中航线对象提供巡检任务的拓扑顺序，航点对象提供经纬度与高度候选，杆塔对象提供近塔暴露与线路走廊约束，基站候选点提供通信节点位置，地形背景提供 DEM/DSM 接入后的高程约束。"
        "这种对象化表达使报告从描述性分析转向结构化分析：每个航点既是飞行任务点，也是通信链路端点、地形剖面端点和 EMI 暴露评价点。"
    )
    add_formula(doc, "Ω = {P_route, P_wp, P_tower, P_bs, G_terrain}")
    add_formula(doc, "z_i = [longitude_i, latitude_i, h_i, d_bs,i, d_tower,i, q_i]")
    add_body_paragraph(
        doc,
        "式中 z_i 表示第 i 个航点的状态向量，q_i 表示该航点相关字段的可验证性等级。"
        "由于当前通信观测量尚未接入，q_i 不能解释为链路质量置信度，而应理解为数据支撑程度：坐标是否有效、高度含义是否明确、是否具备 DEM/DSM 剖面、是否具备通信日志。"
        "因此，本报告的技术边界可以写成数据可验证性函数 η_i，其作用是提示哪些结论属于几何与模型推演，哪些结论需要后续观测量校准。"
    )
    add_formula(doc, "η_i = q_coord,i · q_height,i · q_DEM · q_link")
    add_body_paragraph(
        doc,
        "进一步地，三维仿真沙盘可被视为一个空间知识图谱的原型：航线、航点、杆塔、基站候选点和地形图层并非孤立图层，而是通过空间邻近、拓扑归属和任务阶段形成关联。"
        "例如，一个航点既属于某条 KML 航线，又可能处于某一杆塔邻近域内，同时还与若干基站候选点构成不同的通信路径。"
        "当这些关系被显式表达后，通信可靠性分析就能够从“单点公式计算”扩展为“对象关系推理”，从而更适合描述低空巡检这种空间强约束任务。"
    )
    add_formula(doc, "Omega_sandbox = Route ∪ Waypoint ∪ Tower ∪ Station ∪ TerrainLayer")
    add_formula(doc, "Relation_i = {belongs_to_route, near_tower, linked_station, terrain_profile}")
    add_body_paragraph(
        doc,
        "基于上述对象关系，本文将沙盘模拟给出的空间叠加结果作为模型推演的入口，而不是作为通信质量的终点结论。"
        "也就是说，沙盘可以给出“某航点靠近某杆塔”“某航段可能需要抽取剖面”“某基站候选点与航线存在较大斜距”等空间判断；"
        "但 RSSI、SNR、丢包率、图传码率和真实断链位置仍需由后续测试获得。"
        "这种写法既保留了沙盘模拟对工程分析的支撑作用，也避免把当前尚未观测的数据误写为现场性能结果。"
    )
    add_formula(doc, "R_pre,i = f_model(z_i, Relation_i, theta_assumption)")

    add_table_caption(doc, "表 1 合同要求—本章研究内容—当前实现方式")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("合同要求", "研究内容", "当前实现方式"),
            ("飞行高度影响", "传播路径、FSPL、LoS/NLoS 与 0.6F1 净空。", "公式计算、模型示意、风险分量表达。"),
            ("电力设备 EMI 影响", "电晕、局放、工频场、暂态和金属结构散射。", "基于杆塔邻近性与文献机理建立暴露风险排序。"),
            ("稳定性与可靠性优化", "高度、节点/中继、冗余链路、近塔控制和失联保护。", "建立综合风险模型并给出控制策略。"),
            ("DEM 与三维沙盘", "航线、杆塔、航点、地形剖面和敏感区。", "可视化平台复核空间关系，DEM/DSM 接入后复算净空。"),
        ],
        [Cm(4.0), Cm(6.3), Cm(6.1)],
        header=True,
    )

    add_table_caption(doc, "表 2 数据基础与可验证性边界")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("数据/资源", "当前状态", "可支撑内容", "边界条件"),
            ("KML 航线", f"{len(summary['routes'])} 条航线，{len(summary['waypoints'])} 个航点。", "经纬度、航点高度候选、线路覆盖范围。", "高度基准和轨迹完整性需结合任务记录复核。"),
            ("杆塔台账", f"{len(summary['towers'])} 条记录，{len(summary['valid_towers'])} 条有效经纬度。", "近塔距离、线路走廊和空间邻近性分析。", "杆塔拓扑、导线型号和电压等级待补充。"),
            ("基站候选点", f"{len(summary['stations'])} 个。", "链路距离、站点布设和中继策略初筛。", "缺少天线挂高、增益、发射功率和接收灵敏度。"),
            ("DEM/DSM", "接口预留，数据待接入。", "LoS、0.6F1 净空和敏感区复算。", "需明确来源、分辨率、坐标系和高程基准。"),
            ("通信观测量", "待采集。", "模型校准、阈值修正和质量闭环。", "当前无 RSSI、SNR、丢包率、图传码率、频谱监测和真实断链位置。"),
        ],
        [Cm(2.9), Cm(4.0), Cm(4.8), Cm(4.7)],
        header=True,
    )
    add_body_paragraph(
        doc,
        "如图 1 所示，系统原型的多地图模式与巡检线路覆盖关系已形成统一表达。该图用于说明三条 KML 航线、航点和杆塔对象已经在统一空间框架中叠加，"
        "为后续传播路径、地形剖面和近塔 EMI 暴露分析提供空间索引。"
    )
    add_figure(
        doc,
        figures["web_comparison"],
        "图 1 可视化平台多地图模式与巡检线路覆盖对照，图件类型：系统原型截图。",
        width_in=5.95,
    )

    # Section 2.
    doc.add_page_break()
    doc.add_heading("2. 飞行高度对传播路径、衰减特性与覆盖范围的影响机制", level=1)
    add_body_paragraph(
        doc,
        "飞行高度改变无人机、地面站和地形之间的几何关系，是山区低空巡检通信链路可靠性的核心变量。高度增加通常抬升传播直线，改善视距链路（Line-of-Sight, LoS）条件；"
        "但链路斜距同步增长，自由空间路径损耗随距离对数增加。因此，高度优化不是单向升高，而是在视距改善、0.6F1 净空、路径损耗、巡检成像和飞行安全之间寻求平衡。"
    )
    add_body_paragraph(
        doc,
        "结合三维仿真沙盘的航点高度与空间位置关系，可以把飞行高度影响分解为三类机制。第一类是几何视距机制，即航点高度改变传播直线相对山脊、坡面和塔材结构的位置；"
        "第二类是链路预算机制，即高度改变斜距并影响 FSPL 基础损耗；第三类是业务约束机制，即巡检成像距离、塔体安全距离和飞行控制余量共同限制高度调整幅度。"
        "沙盘在此处承担的是场景化约束功能：它把“高度是否合理”的问题落实到具体航点、具体塔位和具体地形背景中，而不是停留在单一公式层面。"
    )
    add_body_paragraph(
        doc,
        "在山区巡检环境中，低高度航点通常有利于获得较高分辨率的巡检影像，却可能使链路穿越山体背坡、林带或塔材遮挡区；高高度航点能够改善视距和净空，却会增大链路斜距并可能偏离精细巡检任务所需视角。"
        "因此，沙盘模拟所支撑的高度分析应采用分段评价思想：对靠近塔体的短航段，重点关注近塔 EMI、塔材遮挡和飞控安全；对跨越谷地或远离基站的航段，重点关注 LoS、0.6F1 净空和 FSPL 余量。"
        "这种分段思想比全线采用统一高度更符合山区巡检链路的空间非均匀性。"
    )
    add_body_paragraph(
        doc,
        "从链路几何看，高度对传播距离的影响可由三维斜距表达。若 d_h 为地面站到航点的水平距离，h_G 为地面站天线高度，h_UAV 为无人机天线高度，则链路斜距 d_3D 随高度差增加而增长。"
        "高度升高一方面可能使 LoS 穿越山体遮挡的概率下降，另一方面会增大 d_3D 并抬升 FSPL。"
        "因此，沙盘中的航点高度不宜仅按飞行安全或影像角度评价，而应同时进入链路预算方程。"
    )
    add_formula(doc, "d_3D = sqrt(d_h^2 + (h_UAV - h_G)^2)")
    add_formula(doc, "M = P_t + G_t + G_r - L_FSPL - L_extra - S_req")
    add_body_paragraph(
        doc,
        "式中 M 为链路余量，P_t 为发射功率，G_t 与 G_r 为收发天线增益，L_extra 表示地形绕射、多径、植被、雨雾和电磁暴露等附加损耗，S_req 表示业务所需接收门限。"
        "当前缺少设备参数和观测日志，M 不能被赋予现场数值，但该公式说明了高度优化的理论位置：高度通过 d_3D 影响 L_FSPL，通过 LoS 和 Fresnel 条件影响 L_extra，并通过任务约束影响 S_req。"
        "由沙盘模拟给出的航点—基站—杆塔相对关系，正是上述变量进入工程计算前的空间判读基础。"
    )
    add_body_paragraph(
        doc,
        "由图 2 可见，不同高度下传播路径与山体遮挡之间存在显著耦合关系。该图强调两类风险：其一，低高度容易使传播路径转为非视距链路（Non-Line-of-Sight, NLoS）；"
        "其二，即使 LoS 未完全中断，若第一菲涅尔区被山体、树冠或金属构件侵入，仍会产生绕射损耗、多径衰落和链路余量下降。"
    )
    add_figure(
        doc,
        figures["height"],
        "图 2 不同飞行高度下传播路径、山体遮挡和 0.6F1 净空关系，图件类型：模型示意图。",
        width_in=5.55,
    )
    add_body_paragraph(
        doc,
        "在频段与距离分析中，FSPL 可作为链路预算的基础项："
    )
    add_formula(doc, "L_FSPL = 92.45 + 20 log10(f_GHz) + 20 log10(d_km)")
    add_table_caption(doc, "表 3 3 km 与 4 km 下不同频段理论 FSPL")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    rows = [("频段", "3 km (dB)", "4 km (dB)", "说明")]
    for freq, label in [(0.92, "920 MHz"), (2.4, "2.4 GHz"), (5.8, "5.8 GHz")]:
        rows.append((label, f"{fspl_db(freq, 3):.2f}", f"{fspl_db(freq, 4):.2f}", "理论路径损耗基础项。"))
    fill_table(table, rows, [Cm(3.0), Cm(3.0), Cm(3.0), Cm(7.4)], header=True)
    add_body_paragraph(
        doc,
        "如图 3 所示，920 MHz、2.4 GHz 和 5.8 GHz 在 0.5–5 km 范围内的理论路径损耗随距离增加而上升。3–4 km 仅作为工程敏感距离区间或仿真分析区间，"
        "用于提示距离损耗、地形遮挡和电磁暴露因素可能叠加的航段。该区间不应解释为现场性能边界、断链阈值或设备能力验证结果。"
    )
    add_figure(
        doc,
        figures["matlab_fspl"],
        "图 3 不同频段 FSPL 路径损耗曲线，图件类型：理论计算；3–4 km 标注为工程敏感距离区间。",
        width_in=5.55,
    )
    add_body_paragraph(
        doc,
        "如图 4 所示，第一菲涅尔区 F1 与 0.6F1 随距离增加而扩大。低频链路的 FSPL 较低，但波长较长、F1 半径较大，对净空要求更高；"
        "高频链路的 F1 半径较小，但路径损耗和多径敏感性更强。因而频段选择应与地形剖面、航点高度和天线布设条件联合评估。"
    )
    add_body_paragraph(
        doc,
        "由沙盘空间叠加关系进一步推演，FSPL 与 Fresnel 净空并不是相互独立的指标。对于靠近山体侧坡或塔线走廊的航点，频段提高虽然缩小了 F1 半径，却提高了路径损耗和对多径环境的敏感性；"
        "频段降低虽然改善了基础损耗，却要求更大的净空空间。因而在通信策略上，不能简单地把低频等同于更可靠，也不能把高频等同于更高质量图传。"
        "更稳健的做法是将沙盘中的航点位置、杆塔邻近性和基站候选点方位作为先验，再针对不同航段分别评估频段、天线姿态和中继需求。"
    )
    add_formula(doc, "F1(d1,d2) = sqrt(λ d1 d2 / (d1 + d2))")
    add_formula(doc, "C_F = min_x[h_LoS(x) - h_obstacle(x) - 0.6F1(x)]")
    add_body_paragraph(
        doc,
        "其中 λ 为波长，d1 与 d2 分别为剖面点到链路两端的距离，C_F 为全剖面的最小净空裕度。"
        "若 C_F 接近或小于零，则即使传播直线未被完全截断，也应将该航段纳入敏感区复核。"
        "这一结论与沙盘空间判读具有直接对应关系：凡是沙盘中显示航线贴近山体侧坡、穿越塔线密集区域或远离候选基站的航段，均应优先计算 C_F，而不是仅以水平距离作为判断依据。"
    )
    add_figure(
        doc,
        figures["matlab_fresnel"],
        "图 4 第一菲涅尔区 F1 与 0.6F1 随距离变化曲线，图件类型：理论计算。",
        width_in=5.55,
    )

    # Section 3.
    doc.add_page_break()
    doc.add_heading("3. 基于 DEM/DSM 方法框架的山区通信敏感区分析", level=1)
    add_body_paragraph(
        doc,
        "山区通信敏感区应以剖面分析为核心，而不是仅用二维平面距离判定。可视化平台用于定位航线、航点、杆塔和基站候选点的空间关系；"
        "真实 DEM/DSM 接入后，可沿地面站至航点方向采样地形剖面，并计算 LoS、0.6F1 净空和敏感区。"
    )
    add_body_paragraph(
        doc,
        "三维仿真沙盘为敏感区分析提供了从“对象叠加”到“剖面抽取”的中间层。航线在沙盘中的空间形态可以用于识别候选剖面：例如航点密集区、靠近山脊或坡面的航段、与基站候选点距离较大的航段，以及与杆塔金属结构空间邻近的航段。"
        "这些航段并不必然对应通信不稳定，但它们具备进入后续 DEM/DSM 复算和链路预算校核的优先级。"
        "因此，沙盘模拟在方法链条中发挥筛选作用，而 DEM/DSM 和通信日志则用于完成定量判定。"
    )
    add_body_paragraph(
        doc,
        "从学术上看，该流程体现了山区无线传播分析中的尺度耦合问题：宏观尺度上，山体起伏决定 LoS 是否连续；中观尺度上，坡面、林带和线路走廊决定 Fresnel 净空是否充足；"
        "微观尺度上，杆塔、导线和机体姿态改变近场散射与天线有效增益。可视化沙盘能够把这三个尺度组织到同一任务场景中，使风险解释从单一损耗曲线扩展为具有地形、杆塔和航迹约束的空间推理。"
    )
    add_body_paragraph(
        doc,
        "在方法实现上，敏感区识别可分为三步。第一步，在沙盘中选取候选链路对，即地面站或基站候选点与每个航点形成的传播路径；第二步，沿传播路径抽取 DEM/DSM 剖面，建立 h_DEM(x) 或 h_DSM(x)；"
        "第三步，将 LoS 判定、Fresnel 净空和距离敏感区间叠加，形成航点级风险标签。"
        "该方法保留了沙盘的空间直观性，同时避免把在线底图误用为高程数据源。"
    )
    add_formula(doc, "B_LoS = 1, if h_LoS(x) > h_DEM(x), ∀x∈[0,D]; otherwise B_LoS = 0")
    add_formula(doc, "C_min = min_x C(x)")
    add_formula(doc, "R_Fresnel = min(1, max(0, -C_min/C0))")
    add_body_paragraph(
        doc,
        "上述公式中，B_LoS 表示几何视距是否成立，C_min 表示全剖面最小净空余量，C0 为净空风险归一化尺度。"
        "当 B_LoS=0 时，链路进入高遮挡风险；当 B_LoS=1 但 C_min≤0 时，链路可进入中风险或高风险复核；当 B_LoS=1 且 C_min>0 时，仍需结合距离、频段和 EMI 暴露判断综合风险。"
        "因此，沙盘给出的是候选路径与空间对象关系，DEM/DSM 剖面给出的是几何判据，二者共同构成山区通信敏感区识别的理论基础。"
    )
    add_body_paragraph(
        doc,
        "如图 5 所示，三维仿真沙盘能够呈现局部航线、杆塔与地形之间的空间关系。该图用于辅助选择剖面方向和识别需要复核的关键航段，定量遮挡判断仍需由真实 DEM/DSM 数据完成。"
    )
    add_figure(
        doc,
        figures["web_3d"],
        "图 5 三维仿真沙盘中的局部航线、杆塔与地形关系，图件类型：系统原型截图。",
        width_in=5.85,
    )
    add_body_paragraph(
        doc,
        "当前阶段基于 DEM/DSM 方法框架和示意地形剖面进行预评估，真实 DEM 接入后可复算 LoS、0.6F1 净空和敏感区。"
        "剖面判据可写为："
    )
    add_formula(doc, "h_LoS(x) = h_G + x/D · (h_UAV - h_G)")
    add_formula(doc, "C(x) = h_LoS(x) - h_DEM(x) - 0.6F1(x)")
    add_body_paragraph(
        doc,
        "式中 C(x) 表示考虑 0.6F1 后的净空余量。图 6 采用理论示意剖面说明 LoS、地形曲线和净空边界的关系；"
        "当 C(x)>0 时，剖面点满足净空判据；当 C(x)≤0 时，说明 0.6F1 被侵入或 LoS 可能被遮挡，需要复核高度、通信节点或中继方案。"
    )
    add_figure(
        doc,
        figures["matlab_dem"],
        "图 6 理论示意剖面与 0.6F1 净空判据，图件类型：理论示意剖面，未使用真实 DEM。",
        width_in=5.55,
    )
    add_body_paragraph(
        doc,
        "如图 7 所示，航点高度与最小净空余量之间可以建立可解释的函数关系。当前示意参数下尚未满足 0.6F1 净空要求；"
        "曲线随高度增加逐步接近 0 m 阈值，表明高度调整能够改善净空，但仍需结合真实 DEM/DSM 和通信节点参数重新计算。"
        "因此，高度优化应以满足 0.6F1 净空和控制路径损耗为共同目标，而不能仅依赖单一高度增量。"
    )
    add_figure(
        doc,
        figures["matlab_height_margin"],
        "图 7 飞行高度对最小净空余量的影响，图件类型：模型示意图。",
        width_in=5.55,
    )
    add_body_paragraph(
        doc,
        "如图 8 所示，距离与高度两个变量可被投影到统一风险平面，用于表达工程敏感距离区间和高度折中的耦合关系。"
        "图中归一化风险指数限制在 0–1；高风险区域代表模型推演下需要优先复核的组合条件，而非通信质量观测结论。"
    )
    add_body_paragraph(
        doc,
        "该风险平面可与沙盘中的航点属性表建立映射关系：每个航点可根据其到基站候选点的距离、高度候选值、与杆塔的邻近程度和未来 DEM/DSM 净空结果，投影到相应风险区间。"
        "由此，沙盘不再只是三维展示界面，而可扩展为风险图层的承载平台。运维人员可以在沙盘中直接查看风险来源：是距离主导、遮挡主导、净空不足主导，还是近塔 EMI 暴露主导。"
        "这种分因子表达有助于避免把复杂链路问题简化为单一等级，从而提高策略制定的可解释性。"
    )
    add_body_paragraph(
        doc,
        "若将距离、高度、LoS 和 Fresnel 指标统一写成航点级风险向量，则沙盘中的每个航点均可被赋予一组可追溯属性。"
        "这种表达不仅便于地图渲染，也便于后续与通信日志对齐：当某一航点出现链路质量下降时，可以回溯其对应的距离、净空和近塔暴露分量，判断问题是由地形、距离、频段还是电磁环境主导。"
    )
    add_formula(doc, "r_i = [R_distance,i, R_height,i, R_LoS,i, R_Fresnel,i]")
    add_formula(doc, "R_terrain,i = v · r_i^T")
    add_figure(
        doc,
        figures["matlab_risk_heatmap"],
        "图 8 距离—高度二维通信风险热力图，图件类型：工程假设；归一化风险指数范围为 0–1。",
        width_in=5.55,
    )
    add_table_caption(doc, "表 4 LoS 与 0.6F1 净空风险分级")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("等级", "判据", "工程含义", "建议动作"),
            ("低", "LoS 清晰且 C(x)>0。", "几何净空较充分。", "常规执行并保留日志。"),
            ("中", "LoS 未断但局部 C(x)≤0。", "菲涅尔区被侵入。", "复核高度、频段和天线姿态。"),
            ("高", "地形或障碍物遮挡 LoS。", "NLoS 与绕射风险上升。", "调整航点或设置中继。"),
            ("敏感区", "高风险地形叠加 3–4 km 工程敏感距离区间。", "距离损耗与遮挡风险耦合。", "接入 DEM/DSM 后重点复算。"),
        ],
        [Cm(2.0), Cm(4.8), Cm(4.8), Cm(4.8)],
        header=True,
    )

    # Section 4.
    doc.add_page_break()
    doc.add_heading("4. 电力设备 EMI 及气象耦合因素对通信链路的影响", level=1)
    add_body_paragraph(
        doc,
        "电力巡检场景中的 EMI 主要来自高压导线电晕放电、绝缘子和金具局部放电、工频电场/磁场、开关暂态，以及杆塔和导线金属结构造成的反射与散射。"
        "这些因素可通过前门耦合（Front-door Coupling）进入通信或导航接收机，也可通过后门耦合（Back-door Coupling）经机体、线缆、电源或接口影响飞控与任务载荷。"
    )
    add_body_paragraph(
        doc,
        "与普通低空通信场景相比，电力巡检的 EMI 风险具有明显的任务空间相关性。无人机并非随机穿越电磁环境，而是围绕杆塔、导线、绝缘子和金具执行近距离观测，"
        "因此暴露风险与航点到杆塔/导线的最近距离、悬停时间、机体姿态和天线遮挡关系密切相关。三维仿真沙盘能够把这些空间邻近关系显式化，为 R_EMI 的构造提供对象级输入。"
        "这也是本文将 EMI 表述为“电磁暴露风险排序”而非绝对场强估计的原因。"
    )
    add_body_paragraph(
        doc,
        "从耦合机制看，前门耦合主要改变通信链路的等效噪声环境和接收机前端工作状态；后门耦合则可能通过电源、线缆和传感器通道影响导航与飞控系统。"
        "沙盘模拟能够辅助判断无人机是否处于杆塔金属结构包围、导线下方穿越或近塔悬停等空间状态，这些状态是 EMI 暴露排序的重要情景变量。"
        "但由于当前缺少线路电压等级、导线型号、绝缘子状态和频谱监测数据，沙盘只能提供空间暴露线索，不能直接推出干扰强度或链路中断结论。"
    )
    add_body_paragraph(
        doc,
        "理论上，EMI 对链路的影响可被写成接收端信号叠加模型。接收机输入并非只包含期望信号 s(t)，还可能包含经天线进入的前门干扰 i_F(t)、经机体或线缆进入的后门干扰 i_B(t)，以及热噪声和背景噪声 n(t)。"
        "在缺少频谱监测数据时，本文不估计 i_F(t) 与 i_B(t) 的真实幅度，而是通过沙盘中的近塔距离、金属结构邻近性和气象修正项构造暴露排序。"
    )
    add_formula(doc, "y(t) = s(t) + i_F(t) + i_B(t) + n(t)")
    add_formula(doc, "S_dist = exp(-d_min/d0)")
    add_formula(doc, "S_weather = β_r R_rain + β_h H_humidity + β_v V_fog")
    add_body_paragraph(
        doc,
        "式中 d_min 表示航点到杆塔、导线或高压设备候选对象的最近距离，d0 表示暴露衰减尺度；R_rain、H_humidity 与 V_fog 分别表示雨强、湿度和雾/能见度相关修正项。"
        "这些变量可以由沙盘对象关系、气象记录和设备台账逐步补充。"
        "因此，EMI 分析的学术核心不是给出未经观测支持的场强数值，而是建立从空间邻近性到暴露排序、再到链路可靠性控制策略的因果链条。"
    )
    add_body_paragraph(
        doc,
        "如图 9 所示，EMI 风险源、耦合路径和影响对象之间存在链式传导关系。当前 EMI 分析仅为基于杆塔空间邻近性和文献机理的电磁暴露风险排序，不给出真实 EMI 强度、频谱分布、干扰功率或确定性断链结论。"
    )
    add_figure(
        doc,
        figures["emi"],
        "图 9 电力设备 EMI 对无人机通信链路的影响路径，图件类型：模型示意图。",
        width_in=5.55,
    )
    add_formula(doc, "R_EMI = a1S_dist + a2S_tower + a3S_weather + a4S_equipment")
    add_body_paragraph(
        doc,
        "本阶段缺少线路电压等级、导线型号、绝缘子状态、设备运行状态和频谱监测数据，因此 R_EMI 仅表示电磁环境暴露风险，不代表真实干扰场强。"
        "其中 S_dist 可表示近塔距离暴露，S_tower 表示金属结构反射散射风险，S_weather 表示雨、雾、湿度等气象修正项，S_equipment 表示设备类型和运行状态差异。"
    )
    add_body_paragraph(
        doc,
        "如图 10 所示，R_EMI=exp(-d/d0) 可作为表达近塔距离与暴露风险相对关系的工程假设。d0 取 50 m、80 m 和 120 m 时，曲线分别代表不同空间尺度下的衰减假设；"
        "其用途是为近塔安全距离、悬停时间和链路冗余策略提供排序依据，而不是估计真实场强。"
    )
    add_body_paragraph(
        doc,
        "在可视化平台中，R_EMI 可进一步转化为近塔热区或航点属性标签：当航点靠近杆塔、导线或开关设备候选区域时，系统可提高 EMI 暴露权重；当航点远离电力设备且传播路径开阔时，该权重可相应降低。"
        "这种处理方式使 EMI 从抽象机理转化为可在沙盘中显示和复核的空间风险层，便于与飞行高度、链路距离和 DEM/DSM 净空结果共同解释。"
    )
    add_body_paragraph(
        doc,
        "进一步地，R_EMI 可与任务阶段相结合。巡航转场阶段通常以距离和地形遮挡为主导，近塔悬停阶段则更强调 EMI 暴露、金属散射和导航传感器扰动。"
        "若沙盘中某一航点同时具有近塔距离小、悬停时间长、视线穿越塔材结构和高湿天气等条件，则该点应被列为 EMI 复核重点。"
        "这种面向任务阶段的解释方式，比单独讨论电磁干扰源更符合无人机巡检的实际运行逻辑。"
    )
    add_formula(doc, "R_EMI,stage = μ1S_dist + μ2S_tower + μ3S_weather + μ4S_hover")
    add_figure(
        doc,
        figures["matlab_emi_distance"],
        "图 10 近塔距离与 EMI 暴露风险关系，图件类型：工程假设；R_EMI=exp(-d/d0)。",
        width_in=5.55,
    )
    add_table_caption(doc, "表 5 EMI 风险源、影响对象与校准条件")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("风险源", "影响对象", "可能表现", "校准条件"),
            ("电晕/局放", "遥控、图传、数传、GNSS", "噪声底抬升、误码上升。", "频谱监测、天气与缺陷记录。"),
            ("工频场", "磁罗盘、飞控、载荷线缆", "姿态或传感器扰动。", "电压等级、距离和机型抗扰度。"),
            ("开关暂态", "接收机、电源与接口", "短时重传或异常告警。", "设备运行状态记录。"),
            ("金属散射", "图传/数传链路", "多径衰落、吞吐波动。", "频段、天线姿态和空间模型。"),
            ("雨雾湿度", "传播路径与放电活动", "衰减修正或电晕增强。", "气象与链路日志同步采样。"),
        ],
        [Cm(3.0), Cm(3.8), Cm(4.2), Cm(5.4)],
        header=True,
    )

    # Section 5.
    doc.add_page_break()
    doc.add_heading("5. 通信链路优化策略与安全控制建议", level=1)
    add_body_paragraph(
        doc,
        "综合风险模型的目的不是替代现场链路测试，而是将高度、距离、视距、净空、电磁暴露和气象修正拆分为可解释分量，形成任务前复核和后续校准的统一框架："
    )
    add_body_paragraph(
        doc,
        "在系统实现层面，三维仿真沙盘可作为综合风险模型的前端表达层。航线、航点、杆塔、基站候选点和未来 DEM/DSM 剖面结果可被组织为统一对象，"
        "每个对象承载几何距离、高度候选、近塔距离、净空余量和风险等级等属性。这样，优化策略就不再是报告中的静态建议，而可以转化为沙盘中的可视化约束："
        "例如对高风险航段标注调高、绕飞、增设中继或启用备用链路的建议。"
    )
    add_formula(doc, "R_total = w1R_height + w2R_distance + w3R_LoS + w4R_Fresnel + w5R_EMI + w6R_weather")
    add_body_paragraph(
        doc,
        "如图 11 所示，综合风险分级与优化策略可组织为输入、计算、分级和控制四个环节。模型先由航线、杆塔、基站候选点和 DEM/DSM 剖面形成空间风险分量，再叠加 EMI 与天气修正，"
        "最后映射到高度调整、通信节点优化、链路冗余和近塔安全控制等措施。"
    )
    add_body_paragraph(
        doc,
        "从决策逻辑看，综合风险模型应强调“主导因子识别”而不是简单求和。若沙盘显示航段位于山体背向侧且 DEM/DSM 复算后净空不足，则应优先处理 LoS 与 Fresnel 问题；"
        "若航段主要位于开阔区域但距离基站候选点较远，则应优先处理链路预算和中继布设问题；若航点密集分布于杆塔和导线附近，则应优先控制 EMI 暴露、悬停时间和链路冗余。"
        "这种从沙盘空间状态到模型分量再到控制策略的映射，是提高巡检通信可靠性和任务可解释性的关键。"
    )
    add_body_paragraph(
        doc,
        "综合优化可写成多目标约束问题。目标函数既包括通信风险，也包括能耗、任务时间、巡检成像质量和安全裕度。"
        "对低空巡检而言，单纯最小化通信风险可能导致航点高度过高、影像分辨率不足或任务能耗上升；单纯追求成像质量又可能使链路进入遮挡或近塔暴露区。"
        "因此，优化策略应在沙盘中以航点级约束呈现，并由任务需求确定权重。"
    )
    add_formula(doc, "min J = αR_total + βE_flight + γT_task + δQ_image")
    add_formula(doc, "s.t. h_min ≤ h_i ≤ h_max, d_tower,i ≥ d_safe, R_total,i ≤ R_lim")
    add_body_paragraph(
        doc,
        "其中 E_flight 表示能耗代价，T_task 表示任务时间代价，Q_image 表示影像质量偏离代价，d_safe 表示近塔安全距离，R_lim 表示可接受风险阈值。"
        "这些约束均可映射到三维仿真沙盘：高度约束对应航点高度属性，近塔安全距离对应杆塔邻近性，风险阈值对应风险图层颜色和任务控制规则。"
        "由此，优化策略从文字建议转化为可以在系统原型中复核和迭代的决策条件。"
    )
    add_figure(
        doc,
        figures["risk"],
        "图 11 通信链路综合风险分级与优化策略流程，图件类型：工程假设。",
        width_in=5.55,
    )
    add_table_caption(doc, "表 6 综合通信风险分级建议")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("等级", "R_total 区间", "典型条件", "控制建议"),
            ("低", "R < 0.25", "LoS 清晰、0.6F1 满足、距离较短。", "常规执行，保留日志。"),
            ("中", "0.25–0.50", "距离偏大或净空裕度不足。", "复核高度、返航点和备用链路。"),
            ("高", "0.50–0.75", "遮挡、3–4 km 工程敏感距离区间或 EMI 暴露叠加。", "调整节点、抬升航点或设置中继。"),
            ("极高", "R ≥ 0.75", "遮挡、距离、EMI 与恶劣天气多因素耦合。", "重规划任务并开展专项验证。"),
        ],
        [Cm(2.2), Cm(2.6), Cm(6.5), Cm(5.1)],
        header=True,
    )
    add_table_caption(doc, "表 7 风险问题—优化策略—预期效果")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("风险问题", "优化策略", "预期效果"),
            ("地形遮挡或净空不足", "依据 DEM/DSM 剖面调整局部航点高度。", "提高 LoS 概率和 0.6F1 净空余量。"),
            ("距离损耗偏高", "优化地面站/基站位置，必要时设置中继。", "缩短有效链路距离，提高链路预算余量。"),
            ("近塔 EMI 暴露", "控制近塔距离、悬停时间和天线遮挡，配置冗余链路。", "降低短时链路波动和系统扰动风险。"),
            ("导航与飞控受扰", "提高视觉定位、惯导和异常触发规则权重。", "降低定位异常向飞控安全传导。"),
            ("数据闭环不足", "补采 RSSI、SNR、丢包率、图传码率、频谱监测和真实断链位置。", "把预评估模型逐步校准为可验证模型。"),
        ],
        [Cm(4.2), Cm(7.0), Cm(5.2)],
        header=True,
    )
    add_body_paragraph(
        doc,
        "总体而言，优化策略应遵循“先识别主导风险，再选择最小代价控制措施”的原则：地形主导时优先调整高度或节点，距离主导时优先缩短链路或增设中继，"
        "EMI 主导时优先控制近塔暴露和链路冗余，多因素耦合时应重规划任务并开展专项验证。"
    )
    add_body_paragraph(
        doc,
        "后续工程化建议将可视化沙盘、DEM/DSM 数据、通信日志和飞控日志建立闭环：沙盘负责呈现空间对象与风险标签，DEM/DSM 负责复算地形净空，通信日志负责校准链路质量阈值，飞控日志负责校核任务安全状态。"
        "当上述数据逐步完善后，本文提出的风险模型可从预评估工具演进为动态任务规划工具，在任务前给出航线复核，在任务中辅助监测链路状态，在任务后支持故障回放与参数更新。"
    )
    add_body_paragraph(
        doc,
        "最终，通信链路优化不应被理解为单一设备参数优化，而应被理解为“空间路径—传播机制—电磁暴露—任务控制”的协同优化。"
        "三维仿真沙盘提供空间路径的可视化与对象组织，传播模型提供机理解释，风险模型提供排序依据，飞行控制策略提供执行约束。"
        "四者形成闭环后，巡检任务才能在复杂山区和高压电力环境中实现可解释、可复核、可逐步校准的通信可靠性管理。"
    )
    doc.add_heading("参考文献", level=2)
    references = [
        "ITU-R P.525, Calculation of free-space attenuation, International Telecommunication Union.",
        "ITU-R P.526, Propagation by diffraction, International Telecommunication Union.",
        "ITU-R P.530, Propagation data and prediction methods required for the design of terrestrial line-of-sight systems, International Telecommunication Union.",
        "Khawaja, W. et al., A Survey of Air-to-Ground Propagation Channel Modeling for Unmanned Aerial Vehicles, IEEE Communications Surveys & Tutorials.",
        "Khuwaja, A. A. et al., A Survey of Channel Modeling for UAV Communications, IEEE Communications Surveys & Tutorials.",
        "Al-Hourani, A., Kandeepan, S., and Lardner, S., Optimal LAP Altitude for Maximum Coverage, IEEE Wireless Communications Letters.",
        "NASA Earthdata / USGS, Shuttle Radar Topography Mission (SRTM) data documentation.",
        "CISPR 18 series and IEC 61000-4 series, radio interference characteristics of overhead power lines/high-voltage equipment and EMC immunity test methods; detailed edition information requires final manual verification.",
    ]
    for item in references:
        add_reference_item(doc, item)

    doc.save(REPORT_PATH)


def write_notes(summary: dict[str, object]) -> None:
    route_lines = []
    for item in summary["route_stats"]:
        route_lines.append(
            f"- {item['name']}：{item['waypoint_count']} 个航点，KML 高度 {item['min_height']:.3f}–{item['max_height']:.3f} m，"
            f"航线长度候选 {item['length_m']:.2f} m，最近杆塔 {item.get('nearest', 'unknown')}，匹配置信度 {item.get('match_confidence', 'unknown')}。"
        )
    source_lines = []
    for item in summary.get("source_files", []):
        source_lines.append(f"- `{item.get('path')}`，角色：{item.get('role')}，sha256：`{item.get('sha256')}`。")
    notes = f"""# 报告说明：数据来源、模型假设与适用范围

生成日期：2026-06-11

报告文件：`docs/report/泰山低空巡检无人机通信链路可靠性影响机制与优化策略研究.docx`

## 1. 数据来源

本报告使用当前项目仓库内的标准化处理结果，不覆盖原始数据。主要输入如下：

{chr(10).join(source_lines)}

当前结构化统计：

- KML 航线：{len(summary['routes'])} 条。
- KML 航点：{len(summary['waypoints'])} 个。
- 杆塔台账记录：{len(summary['towers'])} 条，其中有效经纬度 {len(summary['valid_towers'])} 条。
- 基站候选点：{len(summary['stations'])} 个。
- 任务记录：{len(summary['tasks'])} 条。
- 可视化平台多地图截图来源：`output/terrain-2d-check.png`、`output/verified-street-map.png`、`output/current-3d-view.png`、`output/map-layer-check-final.png`，分别复制或拼接为 `docs/report/figures/fig_01_web_terrain_imagery_overview.png` 至 `docs/report/figures/fig_05_web_map_modes_comparison.png`。
- MATLAB 仿真图来源：`docs/report/matlab/generate_report_figures.m`，输出目录为 `docs/report/figures_matlab/`。

航线摘要：

{chr(10).join(route_lines)}

## 2. 模型与公式假设

- 自由空间路径损耗 FSPL 使用：`L_FSPL = 92.45 + 20log10(f_GHz) + 20log10(d_km)`。
- 第一菲涅尔区使用常见近似：`F1 = sqrt(lambda*d1*d2/(d1+d2))`，工程净空判据使用 `0.6F1`。
- DEM 剖面判据使用：`h_LoS(x)=h_G+x/D*(h_UAV-h_G)`，`C(x)=h_LoS(x)-h_DEM(x)-0.6F1(x)`。
- EMI 暴露风险使用：`R_EMI = a1S_dist + a2S_tower + a3S_weather + a4S_equipment`。
- 综合风险使用：`R_total = w1R_height + w2R_distance + w3R_LoS + w4R_Fresnel + w5R_EMI + w6R_weather`。

上述模型用于技术预评估、工程复核线索生成和后续测试方案设计；模型参数可通过通信日志、频谱监测和现场链路测试持续校准。

## 3. 适用范围与校准条件

- RSSI、SNR、丢包率、图传码率、频谱监测和断链位置等观测量尚待纳入统一日志。
- 3–4 km 作为工程敏感距离区间或仿真分析区间，用于链路预算、地形遮挡和 EMI 因子联合复核。
- DEM/DSM 接入后可开展真实地形剖面、AGL、净空余量和遮挡条件计算。
- KML 高度字段含义需结合任务说明、起飞基准和设备记录进一步校准。
- Google Earth、Google 3D Tiles、Cesium World Terrain 或在线底图限定为显示底座；地形计算采用可追溯 DEM/DSM 数据源。
- 线路拓扑、杆塔顺序和航线覆盖关系需结合台账、现场资料和人工复核结果确认。
- EMI 风险指标用于电磁暴露排序；干扰场强和设备抗扰度评价需结合频谱监测与标准化测试。

## 4. 参考依据整理

报告整理的公开来源入口和待人工复核依据：

1. ITU-R P.525，Calculation of free-space attenuation，官方入口：https://www.itu.int/rec/R-REC-P.525/en
2. ITU-R P.526，Propagation by diffraction，官方入口：https://www.itu.int/rec/R-REC-P.526/en
3. ITU-R P.530，Propagation data and prediction methods required for the design of terrestrial line-of-sight systems，官方入口：https://www.itu.int/rec/R-REC-P.530/en
4. A. Al-Hourani, S. Kandeepan, S. Lardner, “Optimal LAP Altitude for Maximum Coverage”，IEEE Wireless Communications Letters, 2014，DOI：10.1109/LWC.2014.2342736。
5. W. Khawaja 等，“A Survey of Air-to-Ground Propagation Channel Modeling for Unmanned Aerial Vehicles”，IEEE Communications Surveys & Tutorials, 2019，DOI：10.1109/COMST.2019.2915069。
6. A. A. Khuwaja 等，“A Survey of Channel Modeling for UAV Communications”，IEEE Communications Surveys & Tutorials, 2018，DOI：10.1109/COMST.2018.2856587。
7. NASA Earthdata / SRTM 资料入口：https://www.earthdata.nasa.gov/data/instruments/srtm
8. 电力设备 EMI 相关资料建议人工核验版本：CISPR 18 系列（overhead power lines and high-voltage equipment radio interference characteristics）、IEC 61000-4 系列 EMC 测试与抗扰度方法、EPRI AC Transmission Line Reference Book（corona/radio noise 相关章节）。这些标准/手册多为付费或版本化资料，未在报告中伪造 DOI、页码或引用次数。

## 5. 本次输出文件

- `docs/report/泰山低空巡检无人机通信链路可靠性影响机制与优化策略研究.docx`
- `docs/report/figures/fig_01_web_terrain_imagery_overview.png`
- `docs/report/figures/fig_02_web_topographic_route_overview.png`
- `docs/report/figures/fig_03_web_3d_route_detail.png`
- `docs/report/figures/fig_04_web_basic_grid_mode.png`
- `docs/report/figures/fig_05_web_map_modes_comparison.png`
- `docs/report/figures/fig_06_height_path_fresnel.png`
- `docs/report/figures/fig_07_dem_profile_fresnel_clearance.png`
- `docs/report/figures/fig_08_fspl_curve_sensitive_distance.png`
- `docs/report/figures/fig_09_emi_coupling_paths.png`
- `docs/report/figures/fig_10_integrated_risk_workflow.png`
- `docs/report/matlab/generate_report_figures.m`
- `docs/report/figures_matlab/matlab_fig_01_fspl_curve.png`
- `docs/report/figures_matlab/matlab_fig_02_fresnel_radius.png`
- `docs/report/figures_matlab/matlab_fig_03_dem_fresnel_clearance.png`
- `docs/report/figures_matlab/matlab_fig_04_height_clearance_margin.png`
- `docs/report/figures_matlab/matlab_fig_05_distance_height_risk_heatmap.png`
- `docs/report/figures_matlab/matlab_fig_06_emi_distance_risk.png`
- `docs/report/report_notes.md`

## 6. 后续需要补充的数据

- DEM/DSM，需说明数据源、分辨率、坐标系、高程基准和许可范围。
- 无人机和地面站/基站通信参数：频段、带宽、发射功率、接收灵敏度、天线增益、天线挂高、方位角、下倾角。
- 飞行日志和通信日志：RSSI、SNR、丢包率、图传码率、断链/重连时间位置、GNSS 质量、飞控告警。
- 电力设备运行状态和电磁环境记录：电压等级、负荷状态、开关操作、局放/电晕观测、频谱监测。
- 航线与任务、杆塔与线路拓扑的人工复核结果。
"""
    NOTES_PATH.write_text(notes, encoding="utf-8")


def write_revision_notes() -> None:
    notes = """# 修订说明 revision_notes

生成日期：2026-06-11

## 1. 篇幅与结构压缩

- 将正文组织为 5 个主节：研究任务与技术边界、飞行高度影响机制、DEM/菲涅尔区通信敏感区、EMI 影响、通信链路优化策略。
- 将原先重复出现的 FSPL、Fresnel、DEM 剖面和风险指数解释压缩为“公式—图件—工程含义”的单次论证链。
- 将插图由 14 张压缩为 11 张核心图，保留系统原型截图、传播机理图、理论计算图、模型示意图和工程假设图。
- 删除正文末尾多余展开段落，避免由于大图和连续分页造成空白页风险。
- 在压缩重复内容的基础上，重新扩展五个主节的学术化论述，重点增加三维仿真沙盘对空间对象叠加、航点高度判读、DEM/DSM 剖面筛选、近塔 EMI 暴露排序和综合风险映射的支撑作用。
- 沙盘相关表述统一为“空间叠加关系”“沙盘模拟框架”“可视化平台风险层”和“模型推演”，不将未由真实 DEM/DSM 或通信观测量产生的内容写成确定性实测结论。
- 按“每章约 2000 字”的目标进一步扩展理论论证，并新增对象集合、航点状态向量、数据可验证性、三维斜距、链路余量、Fresnel 半径、LoS 指示函数、净空风险、EMI 耦合信号模型和多目标优化约束等公式。

## 2. 图号、图注与正文引用修正

- 图号统一为图 1 至图 11，正文在插图出现前先引用对应图号。
- 修正原图 8、图 9、图 10、图 11 附近容易出现的图文错位：修订后图 6 为理论示意剖面，图 7 为高度—净空余量，图 8 为距离—高度风险热力图，图 9 为 EMI 影响路径。
- 图题均位于图下；表题均位于表上。
- 图注均明确标注图件类型：系统原型截图、理论计算、理论示意剖面、模型示意图或工程假设。

## 3. 理论计算图与模型示意图边界

- 理论计算图：图 3 FSPL 路径损耗曲线；图 4 第一菲涅尔区 F1 与 0.6F1 曲线。
- 模型示意图：图 2 飞行高度与遮挡机制；图 7 飞行高度与最小净空余量；图 9 EMI 影响路径。
- 理论示意剖面：图 6 采用示意地形剖面说明 LoS 与 0.6F1 净空判据，未使用真实 DEM。
- 工程假设图：图 8 距离—高度二维通信风险热力图；图 10 近塔距离与 R_EMI 暴露风险；图 11 综合风险分级流程。
- 系统原型截图：图 1 多地图模式与线路覆盖；图 5 三维仿真沙盘局部空间关系。

## 4. MATLAB 图质量修正

- FSPL 图例改为仅显示 920 MHz、2.4 GHz、5.8 GHz 三条频段曲线，避免 data1、data2 等残留项。
- MATLAB 图标题去除括号及括号内说明，理论计算、模型示意、工程假设信息改由图内注释与 Word 图注表达。
- 风险热力图的归一化风险指数已裁剪到 0–1。
- 飞行高度—最小净空余量图在当前示意参数下全程低于 y=0，正文已改为“当前示意参数下尚未满足 0.6F1 净空要求”。
- 所有图维持白底、浅灰网格、清晰坐标单位和规范图例。

## 5. 不可作为实测结论的内容

- 当前无 RSSI、SNR、丢包率、图传码率、频谱监测和真实断链位置，报告中的风险指数用于任务前预评估与后续校准设计。
- 3–4 km 仅作为工程敏感距离区间或仿真分析区间，不作为现场性能边界、断链阈值或设备能力验证结果。
- R_EMI 仅表示基于杆塔空间邻近性和文献机理的电磁暴露风险排序，不代表真实干扰场强、频谱分布或干扰功率。
- 当前 DEM/DSM 数据待接入，报告中的剖面图为理论示意剖面或模型示意图，不能作为真实 DEM 遮挡计算结论。

## 6. 参考文献核验

- 报告已将“参考依据摘要与复核提示”改为“参考文献”列表。
- ITU-R P.525、P.526、P.530、UAV A2G 信道综述、UAV 信道建模综述、Al-Hourani 最优高度模型、SRTM/NASA Earthdata 或 USGS SRTM、CISPR/IEC 电磁兼容资料已列入参考文献。
- 由于当前未进行联网逐条核验，论文卷期页码、DOI、标准版本和付费标准条款需人工最终核验；本文未伪造 DOI、页码或引用次数。
"""
    REVISION_NOTES_PATH.write_text(notes, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    summary = summarize_data()
    figures = make_report_figures(summary)
    build_docx_revised(summary, figures)
    write_notes(summary)
    write_revision_notes()
    print(f"generated: {REPORT_PATH}")
    print(f"generated: {NOTES_PATH}")
    print(f"generated: {REVISION_NOTES_PATH}")
    print(f"figures: {FIG_DIR}")


if __name__ == "__main__":
    main()
