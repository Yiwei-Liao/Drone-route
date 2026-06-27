"""Build revised patent disclosure DOCX files for the Taishan drone project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "docs" / "patents"

BODY_FONT = "SimSun"
LATIN_FONT = "Calibri"
HEADING_COLOR = RGBColor(31, 78, 121)
MUTED_COLOR = RGBColor(90, 90, 90)
TABLE_HEADER_FILL = "F2F4F7"
LIGHT_FILL = "F8F9FB"


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None) -> None:
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_cell_margins(table, top: int = 80, start: int = 120,
                           bottom: int = 80, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
    set_table_cell_margins(table)


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = HEADING_COLOR
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = ""
    run = header.add_run("核心专利交底书")
    set_run_font(run, size=9, color=MUTED_COLOR)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run("核心专利交底书")
    set_run_font(run, size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    set_run_font(run, size=15, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=10.5, color=MUTED_COLOR)
    return doc


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=MUTED_COLOR)
    doc.add_paragraph()


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_table(doc: Document, rows: list[list[str]], widths: list[int], header: bool = False) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = table.cell(ri, ci)
            if header and ri == 0:
                set_cell_shading(cell, TABLE_HEADER_FILL)
            elif ci == 0 and len(rows[0]) == 2:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=10.5, bold=(header and ri == 0) or (ci == 0 and len(rows[0]) == 2))
    doc.add_paragraph()


def add_application_info(doc: Document, patent_name: str, tech_summary: str, prospect: str) -> None:
    add_heading(doc, "第一部分 申请信息", 1)
    rows = [
        ["专利名称*（发明人填写）", patent_name],
        ["核心说明*：技术情况", tech_summary],
        ["核心说明*：运用前景", prospect],
        ["项目类型*", "【待补充】"],
        [
            "依托项目*",
            "产出该专利的项目名称：【待补充】\n"
            "产出该专利的项目编号：【待补充】\n"
            "项目任务书要求产出的专利数量：【待补充】\n"
            "项目牵头单位：【待补充】",
        ],
        ["专利申请类型（发明人填写）*", "发明；是否同步申请实用新型：【待补充】"],
        ["申请人（两个以上申请人请按先后顺序填写）*", "【待补充】"],
        ["发明人（两个以上发明人请按先后顺序填写）*", "【待补充】"],
        ["第一发明人身份证号*", "【待补充】"],
        ["技术联系人*", "姓名：【待补充】；电话：【待补充】；邮箱：【待补充】"],
    ]
    add_table(doc, rows, [2300, 7060])


def add_common_preface(doc: Document) -> None:
    add_heading(doc, "第二部分 技术内容", 1)
    add_note(
        doc,
        "说明：本文档用于向专利代理人交底技术方案、实施方式和可保护要点。"
        "未明确给出的实验数值、设备型号、精度指标、人员信息、身份证号、项目编号、正式阈值、权重和公开文献检索结论均以【待补充】或可配置参数表述。"
        "本文不主张公知公式、普通 GIS 输出格式、CSV/GeoJSON/API 载体本身作为核心创新。"
    )


def add_terms_one(doc: Document) -> None:
    add_heading(doc, "一、缩略语和关键术语定义", 2)
    add_para(
        doc,
        "KML（Keyhole Markup Language，地理标记语言）：用于表达无人机航点、航线或空间对象的数据文件。"
        "Point：KML 中的点几何对象。LineString：KML 中的线几何对象。DJI/MIS 扩展字段：KML 中与 speed（速度）、heading（航向）、gimbalPitch（云台俯仰角）、turnMode（转弯模式）等相关的扩展标签。"
        "杆塔台账：记录线路名称、杆号、经度、纬度、杆塔全高、档距、机场站点等信息的基础台账。"
        "航线半径 R_route：航点集合相对于航线中心的最大空间距离，用于判断航线是否围绕局部杆塔展开。"
        "匹配置信度：基于航点到杆塔距离、航线半径、航点数量和可配置阈值输出的 high、medium 或 low 状态。"
    )


def add_terms_two(doc: Document) -> None:
    add_heading(doc, "一、缩略语和关键术语定义", 2)
    add_para(
        doc,
        "KML（Keyhole Markup Language，地理标记语言）：用于表达无人机航点或航线的数据文件。"
        "DEM（Digital Elevation Model，数字高程模型）：表达地形地面高程的数据。DSM（Digital Surface Model，数字表面模型）：表达地表含建筑、树冠等表面高程的数据。"
        "AGL（Above Ground Level，离地高度）：无人机或航点相对地形或地表的高度。"
        "LoS（Line of Sight，视距）：通信源与航点之间几何视线未被遮挡的状态。NLoS（Non-Line of Sight，非视距）：通信源与航点之间几何视线被遮挡的状态。"
        "FSPL（Free Space Path Loss，自由空间路径损耗）：用于估计自由空间传播损耗的公知模型。Fresnel（菲涅尔）区：无线传播路径周围影响传播能量的空间区域。"
        "EMI（Electromagnetic Interference，电磁干扰）：电力设备或环境因素可能对无人机通信、导航或传感器产生的电磁扰动。"
    )


def add_prior_art_search(doc: Document, directions: list[str]) -> None:
    add_heading(doc, "八、待检索的现有技术方向", 2)
    add_para(doc, "以下方向用于后续专利代理人或检索人员开展人工检索核验，本文不替代正式检索报告，也不伪造 DOI、授权状态、公开日或对比文件。")
    add_bullets(doc, [f"{item}：待人工检索核验。" for item in directions])


def add_figure_descriptions_one(doc: Document) -> None:
    add_heading(doc, "3.2.7 附图说明", 3)
    add_bullets(doc, [
        "附图1为多源巡检数据标准化流程图，示出原始巡检任务表、线路台账、KML 航线文件只读输入后，经过字段识别、KML 解析、航线指标计算、杆塔匹配、质量检查和结果输出的流程。",
        "附图2为系统模块结构示意图，示出数据源管理模块、字段识别与映射模块、KML 解析模块、航线指标计算模块、杆塔匹配模块、数据质量闭环模块和结果发布模块之间的数据流。",
        "附图3为 KML 航线与杆塔匹配示意图，示出航点序列、航线中心、航线半径 R_route、候选杆塔、最小距离 d_min、平均距离 d_avg 和匹配置信度输出。",
        "附图4为数据质量闭环示意图，示出缺失坐标、疑似经纬度反置、重复 ID 和低置信匹配如何形成质量问题、补录模板和补录校验结果。",
    ])


def add_figure_descriptions_two(doc: Document) -> None:
    add_heading(doc, "3.2.7 附图说明", 3)
    add_bullets(doc, [
        "附图1为航点级地形净空与通信链路联合评估流程图，示出航线对象、通信源、DEM/DSM、剖面采样、LoS 判定、Fresnel 净空判定和风险标签输出。",
        "附图2为高度语义统一示意图，示出 KML 高度、航线设定高度、地形高程、杆塔高度、展示高度和真实 AGL 之间的区别。",
        "附图3为通信源至航点的剖面采样示意图，示出采样点 Q_m、链路线高度 h_LoS、地形高程 h_DEM、Fresnel 半径 F1 和净空余量 C。",
        "附图4为状态机示意图，示出 terrain pending、distance-only、parameter missing、LoS clear、Fresnel risk 和 NLoS blocked 等状态及其转换条件。",
        "附图5为 GIS 风险图层输出示意图，示出风险航段、风险点、通信走廊、剖面证据图和待补充数据清单。"
    ])


def build_disclosure_one() -> None:
    name = "一种面向电力无人机巡检的多源数据标准化、航线解析与杆塔匹配方法及系统"
    doc = setup_document(name, "项目交底书一：多源数据标准化、KML 航线解析与杆塔匹配")
    doc.core_properties.title = name
    doc.core_properties.author = "【待补充】"

    add_application_info(
        doc,
        name,
        "本发明面向电力无人机巡检中原始 Excel/KML 数据来源分散、中文字段语义不一致、航线与杆塔空间对象容易错配且处理过程不可追溯的问题，提出一种多源数据标准化、航线解析与杆塔匹配方法及系统。"
        "该系统对原始巡检 Excel/KML 文件进行只读清点并记录路径、角色、大小、修改时间和哈希值；对中文字段、复合表头、重复字段和合并单元格语义进行标准映射；对 KML Point 优先解析并在必要时回退 LineString 解析；将扩展字段按航点序号对齐；并基于航点到杆塔距离、航线半径、航点数量和可配置阈值生成杆塔匹配置信度及匹配理由。",
        "本发明可用于无人机巡检数据治理、线路台账复核、航线与杆塔对应关系确认、地图可视化前的数据可信处理以及后续地形/通信评估的数据输入。"
        "其技术边界在于：不覆盖原始数据，不自动补造缺失坐标，不将待确认字段写成确定结论，不把普通 CSV、GeoJSON、SQLite 或 API 输出载体本身作为核心创新。",
    )
    add_common_preface(doc)
    add_terms_one(doc)

    add_heading(doc, "二、相关技术背景（背景技术），与本发明最相近似的现有实现方案（现有技术）", 2)
    add_heading(doc, "2.1 技术领域", 3)
    add_para(doc, "本发明属于电力无人机巡检数据处理、KML 航线解析、线路杆塔台账 GIS 匹配和空间数据质量控制技术领域。")
    add_heading(doc, "2.2 背景技术", 3)
    add_para(doc, "电力无人机巡检任务通常同时产生巡检任务统计表、线路杆塔台账、无人机 KML 航线文件和地图展示数据。不同来源数据在字段命名、表头结构、坐标可用性、高度语义和文件组织方式上存在明显差异。若缺少统一的数据标准化和质量状态记录，KML 航线可能被错误地解释为完整巡线轨迹，缺失坐标杆塔可能被人为推测为地图点，航线与杆塔之间的最近邻关系也可能缺少可追溯的判断依据。")
    add_para(doc, "在一种实施方式中，原始任务表可包括复合表头、重复时间字段、气象字段、飞行距离、飞行时长、电池电量、基站编号和基站经纬度；线路台账可包括线路名称、杆号、经度、纬度、杆塔全高、档距和机场站点；KML 航线可包括航点经纬度、高度以及 speed、heading、gimbalPitch、turnMode 等扩展字段。上述字段若直接用于地图显示，容易形成空间对象错配和不可复核结果。")
    add_heading(doc, "2.3 现有技术一及其缺点", 3)
    add_para(doc, "现有技术一通常由人工分别打开 Excel 台账和 KML 文件，人工识别字段并复制到 GIS 软件或电子表格中进行叠加显示。该方式的缺点在于：字段映射依赖人工经验，复合表头、重复字段和合并单元格语义难以稳定保留；原始文件路径、哈希值、源行号和字段含义无法形成自动追溯；后续发现坐标缺失或疑似反置时难以定位原始记录。")
    add_heading(doc, "2.4 现有技术二及其缺点", 3)
    add_para(doc, "现有技术二通常直接解析 KML 航线并与杆塔点图层做最近邻或视觉判断。该方式虽然能得到直观叠加效果，但常常未区分 KML Point 和 LineString 的优先级，未对 DJI/MIS 扩展字段按航点序号对齐，未结合航线半径和航点数量解释航线是否为单塔或局部巡检，且缺少 high、medium、low 等置信度和匹配理由输出。")
    add_heading(doc, "2.5 核心区别特征", 3)
    add_numbered(doc, [
        "原始巡检 Excel/KML 文件只读清点，记录文件路径、角色、大小、修改时间和哈希值，使数据处理链路可追溯。",
        "面向中文字段、复合表头、重复字段和合并单元格语义建立字段识别与标准映射，不把无法确认含义的字段写成已解析字段。",
        "KML 解析采用 Point 优先、LineString 回退机制，并将 speed、heading、gimbalPitch、turnMode 等扩展字段按航点序号对齐。",
        "基于航点到杆塔距离、航线半径、航点数量和可配置阈值生成匹配置信度和匹配理由，而非仅输出最近杆塔。",
        "对缺失坐标、疑似经纬度反置、重复 ID、低置信匹配输出质量问题、补录模板和补录校验结果，且不自动补造坐标。"
    ])

    add_heading(doc, "三、本发明技术方案的详细阐述（技术方案）", 2)
    add_heading(doc, "3.1 本发明所要解决的技术问题（发明目的）", 3)
    add_para(doc, "本发明解决的技术问题不是普通格式转换，而是解决电力无人机巡检中多源空间对象的可信建模问题：其一，解决原始巡检 Excel/KML 数据在处理前缺少只读清点和文件级追溯的问题；其二，解决中文字段、复合表头、重复字段和合并单元格语义导致标准字段难以稳定生成的问题；其三，解决 KML 航点及扩展字段无法按航点序列形成完整航线对象的问题；其四，解决航线与杆塔匹配缺少置信度、匹配理由和质量闭环的问题。")
    add_heading(doc, "3.2 系统组成", 3)
    add_numbered(doc, [
        "数据源清点模块，用于对原始巡检 Excel 文件和 KML 文件进行只读扫描，记录文件路径、文件角色、文件大小、修改时间、哈希值以及扫描目录状态。",
        "字段识别与标准映射模块，用于识别中文字段、复合表头、重复字段、合并单元格语义和单位待确认字段，并生成原始字段到标准字段的映射关系。",
        "KML 航线解析模块，用于解析 KML Point 或 LineString 坐标以及 speed、heading、gimbalPitch、turnMode 等扩展字段。",
        "航线指标计算模块，用于计算航点数量、航线长度、航点间距、高度范围、航线中心和航线半径 R_route。",
        "杆塔匹配模块，用于计算航点到候选杆塔的最小距离 d_min、平均距离 d_avg、最大距离 d_max，并输出匹配置信度和匹配理由。",
        "数据质量闭环模块，用于识别缺失坐标、疑似经纬度反置、重复 ID 和低置信匹配，输出质量问题、补录模板和补录校验结果。",
        "结果发布模块，用于将标准表和空间图层发布给后续仿真、GIS 展示或接口调用；所述 CSV、GeoJSON、SQLite 和 API 仅为输出载体。"
    ])

    add_heading(doc, "3.3 方法流程", 3)
    add_numbered(doc, [
        "步骤S1，读取并清点原始文件，禁止覆盖或修改原始文件；为每个文件生成文件角色、路径、大小、修改时间和哈希值。",
        "步骤S2，识别巡检任务表和线路台账中的中文字段、复合表头、重复字段和合并单元格语义，建立标准字段映射表。",
        "步骤S3，生成杆塔、线路、任务、基站和机场候选标准对象；其中经纬度字段统一为 longitude 和 latitude，无法确认含义的高度或距离字段标记为 unknown 或待确认。",
        "步骤S4，解析 KML 文件，优先提取 Point 航点；若 Point 航点不存在或为空，则回退提取 LineString 坐标；对扩展字段按航点序号进行对齐，不足项为空。",
        "步骤S5，基于航点序列计算航线几何指标，包括航线长度、航点间距、高度范围、航线中心和航线半径 R_route。",
        "步骤S6，针对每条航线与有效坐标杆塔计算距离统计，生成最近杆塔、d_min、d_avg、d_max、匹配置信度和匹配理由。",
        "步骤S7，执行数据质量检查，输出缺失坐标、疑似经纬度反置、重复 ID 和低置信匹配等质量状态。",
        "步骤S8，针对缺失坐标问题生成补录模板和补录校验结果；补录模板仅用于人工回到可信来源核验，不自动回写原始台账。",
        "步骤S9，输出标准对象表、质量问题表和可视化图层，供后续系统读取。"
    ])

    add_heading(doc, "3.4 算法化表达", 3)
    add_para(doc, "设某条航线包含 N 个航点 P_i=(lon_i,lat_i,h_i)，i=1,...,N。系统可将经纬度转换到局部平面坐标或使用大地距离函数 dist(·,·) 进行距离计算。航线中心 C_route 可由航点坐标的几何中心、加权中心或局部投影中心得到。航线半径定义为：R_route=max_i dist(P_i,C_route)。")
    add_para(doc, "设候选杆塔集合为 T_j，j=1,...,M。对于航线 r 与杆塔 T_j，航点到杆塔的距离集合为 D_j={dist(P_i,T_j)|i=1,...,N}。最小距离 d_min(j)=min(D_j)，平均距离 d_avg(j)=sum(D_j)/N，最大距离 d_max(j)=max(D_j)。系统选择 d_min 最小的杆塔作为最近候选杆塔，并保留 d_min、d_avg、d_max 用于置信度判定。")
    add_para(doc, "匹配阈值、航线半径阈值和航点数量阈值均为可配置参数，可根据地区、线路类型和业务要求进行工程标定，不在本文中写死。")
    add_code_block(doc, [
        "for each route r:",
        "    if r has no valid waypoint or no valid tower coordinate:",
        "        confidence = low; reason = missing spatial data",
        "    else:",
        "        compute R_route, N, d_min, d_avg, d_max for nearest tower",
        "        if d_min <= TH_HIGH_DISTANCE and R_route <= TH_ROUTE_RADIUS and N <= TH_WAYPOINT_COUNT:",
        "            confidence = high",
        "            reason = route is close to tower and spatially compact",
        "        elif d_min <= TH_MEDIUM_DISTANCE:",
        "            confidence = medium",
        "            reason = route is near tower but needs manual confirmation",
        "        else:",
        "            confidence = low",
        "            reason = route is far from nearest tower or matching condition is insufficient",
    ])
    add_para(doc, "数据质量状态规则可表示如下：若 longitude 或 latitude 缺失、为空或无法转为数值，则输出 missing_coordinate；若 longitude 落入配置的纬度经验范围且 latitude 落入配置的经度经验范围，则输出 possible_lon_lat_swapped；若标准 ID 在同一标准表中重复，则输出 duplicate_id；若匹配置信度为 low，则输出 low_match_confidence；若补录坐标只填写一列、无法转数值、疑似反置或超出配置经验范围，则分别输出 incomplete、invalid_numeric、possible_lon_lat_swapped 或 out_of_range。")

    add_heading(doc, "3.5 输出标准表字段示例", 3)
    add_table(doc, [
        ["标准表", "字段示例", "说明"],
        ["routes", "route_id, route_name, kml_file, geometry_used, waypoint_count, total_length, R_route, min_height, max_height, route_type_guess", "记录航线对象和航线几何指标；CSV 等仅为载体。"],
        ["route_waypoints", "waypoint_id, route_id, sequence, longitude, latitude, altitude, speed, heading, gimbal_pitch, turn_mode", "记录航点序列和按序号对齐的扩展字段。"],
        ["route_tower_matches", "route_id, tower_id, d_min, d_avg, d_max, match_confidence, match_reason", "记录航线与杆塔匹配结果及可解释依据。"],
        ["data_quality_issues", "table, entity_id, severity, issue_type, message, source_file, source_row, longitude, latitude", "记录质量问题和回到原始记录复核所需的上下文。"],
    ], [1700, 4450, 3210], header=True)

    add_heading(doc, "3.6 实施例", 3)
    add_para(doc, "在一种实施方式中，系统读取巡检任务统计表、线路台账和多个 KML 航线文件。系统首先生成文件清单和哈希记录；然后将线路台账中杆塔经纬度标准化为 longitude 和 latitude，将任务表中基站经纬度拆分为 station_longitude 和 station_latitude；再解析 KML 航点和扩展字段，形成航线对象和航点对象。")
    add_para(doc, "在另一种实施方式中，当线路台账中部分杆塔缺少经纬度时，系统将该类记录写入 data_quality_issues，并生成坐标补录模板。该模板中的待填经度和待填纬度保持空白，需人工依据可信台账、现场资料或测绘数据补充；系统仅校验补录值的完整性、数值性和经验范围，不自动生成或回写坐标。")
    add_figure_descriptions_one(doc)

    add_heading(doc, "四、技术方案产生的效果（本发明技术方案带来的有益效果）", 2)
    add_numbered(doc, [
        "通过只读清点和哈希记录，提升原始巡检数据处理链路的可追溯性。",
        "通过中文字段、复合表头、重复字段和合并单元格语义处理，降低标准字段错配风险。",
        "通过 Point 优先和 LineString 回退解析，提高 KML 航点提取的兼容性。",
        "通过扩展字段按航点序号对齐，使速度、航向、云台俯仰角和转弯模式能够与航点空间位置关联。",
        "通过距离、航线半径和航点数量联合生成匹配置信度，避免仅按最近点强制确认杆塔对象。",
        "通过质量问题、补录模板和补录校验结果，形成不伪造坐标的数据整改闭环。"
    ])

    add_heading(doc, "五、可替代方案", 2)
    add_para(doc, "KML 解析模块可采用 XML 解析器、GIS 库或无人机厂商 SDK 实现；距离计算可采用大地距离、局部投影距离、空间索引最近邻或线段到点距离；字段映射规则可采用配置文件、字典表或交互式确认方式。上述替代方案不改变本发明的核心，即在可追溯清点、标准映射、航线解析、置信匹配和质量闭环之间形成组合技术方案。")

    add_heading(doc, "六、技术关键点和欲保护点", 2)
    add_numbered(doc, [
        "保护原始巡检 Excel/KML 文件只读清点并记录路径、角色、大小、修改时间和哈希值的处理机制。",
        "保护面向中文字段、复合表头、重复字段和合并单元格语义的标准字段映射机制。",
        "保护 KML Point 优先、LineString 回退以及扩展字段按航点序号对齐的航线解析机制。",
        "保护基于航点到杆塔距离、航线半径、航点数量和可配置阈值生成匹配置信度及匹配理由的机制。",
        "保护缺失坐标、疑似经纬度反置、重复 ID 和低置信匹配的质量问题输出、补录模板生成和补录校验机制。",
        "保护将上述标准化和匹配结果作为后续仿真或可视化输入的组合流程；CSV、GeoJSON、SQLite、API 等仅作为载体。"
    ])

    add_heading(doc, "七、权利要求草案", 2)
    add_heading(doc, "7.1 独立权利要求草案", 3)
    add_numbered(doc, [
        "一种多源巡检数据标准化与杆塔航线匹配系统，包括：数据源清点模块、字段识别与标准映射模块、KML 航线解析模块、航线指标计算模块、杆塔匹配模块、数据质量闭环模块和结果发布模块；其中，数据源清点模块用于对原始巡检 Excel/KML 文件进行只读清点并记录文件路径、角色、大小、修改时间和哈希值；字段识别与标准映射模块用于处理中文字段、复合表头、重复字段和合并单元格语义；KML 航线解析模块用于优先解析 Point 航点并在 Point 不可用时回退解析 LineString 坐标，且将扩展字段按航点序号对齐；杆塔匹配模块用于基于航点到杆塔距离、航线半径、航点数量和可配置阈值生成匹配置信度和匹配理由；数据质量闭环模块用于输出质量问题、补录模板和补录校验结果，并禁止自动补造坐标。",
        "一种多源巡检数据标准化与杆塔航线匹配方法，包括：只读清点原始巡检 Excel/KML 文件；识别并映射中文字段、复合表头、重复字段和合并单元格语义；解析 KML 航点及扩展字段；计算航线半径和航点到杆塔的距离统计；根据可配置阈值输出杆塔匹配置信度和匹配理由；对缺失坐标、疑似经纬度反置、重复 ID 和低置信匹配生成质量问题、补录模板和补录校验结果。",
        "一种电子设备或计算机可读存储介质，其上存储有计算机程序，所述计算机程序被处理器执行时实现上述多源巡检数据标准化与杆塔航线匹配方法。"
    ])
    add_heading(doc, "7.2 从属权利要求要点", 3)
    add_bullets(doc, [
        "字段映射进一步包括对原始字段、标准字段、单位、来源文件和源行号的关联记录。",
        "KML 回退解析进一步包括在 Point 航点为空时读取 LineString 坐标，并记录 geometry_used 字段。",
        "扩展字段对齐进一步包括将 speed、heading、gimbalPitch、turnMode 按航点 sequence 对齐，缺失项为空。",
        "航线指标计算进一步包括航线长度、平均航点间距、最大航点间距、高度范围和 R_route。",
        "匹配置信度进一步包括 high、medium、low 三类状态及对应 match_reason。",
        "质量闭环进一步包括 coordinate_backfill_template 和 coordinate_backfill_validation。",
        "结果发布进一步包括将标准对象供 GIS 或接口读取，但不以 CSV、GeoJSON 或 API 载体本身作为核心创新。"
    ])

    add_prior_art_search(doc, [
        "KML 航线解析",
        "无人机巡检数据治理",
        "杆塔台账 GIS 匹配",
        "空间最近邻匹配",
        "数据质量闭环"
    ])
    add_heading(doc, "九、参考文献（如专利/论文/标准等）", 2)
    add_para(doc, "公开专利、论文、标准和 DOI 信息均待人工检索核验；本文不伪造检索结论。")
    save_doc(doc, OUT_DIR / "专利交底书1_多源巡检数据标准化与杆塔航线匹配系统及方法.docx")


def build_disclosure_two() -> None:
    name = "一种面向山区电力无人机巡检的航点级地形净空与通信链路风险联合评估方法及系统"
    doc = setup_document(name, "项目交底书二：航点级地形净空与通信链路风险联合评估")
    doc.core_properties.title = name
    doc.core_properties.author = "【待补充】"

    add_application_info(
        doc,
        name,
        "本发明面向山区电力无人机巡检中航点高度、地形净空、通信源距离、Fresnel 净空和近塔电磁暴露候选难以统一评估的问题，提出一种航点级地形净空与通信链路风险联合评估方法及系统。"
        "该系统以航点级对象统一处理航线、杆塔、通信源和 DEM/DSM 地形数据，区分 KML 高度、航线设定高度、地形高程、杆塔高度、展示高度和真实 AGL；当 DEM/DSM 不可用时输出 DEM not provided 或 terrain pending，不计算真实 AGL、真实遮挡或工程覆盖结论；并在通信源到航点的剖面采样链路上联合计算 LoS、可配置比例的第一 Fresnel 区净空、距离链路压力和近塔 EMI 暴露候选。",
        "本发明可用于山区电力无人机巡检的航点级或航段级风险标签生成、GIS 风险图层输出、剖面证据图生成和待补充数据清单整理。"
        "通信设备参数、DEM/DSM 分辨率、风险权重和阈值均为可配置或待工程标定参数；本文不写入实测 RSSI、SNR、丢包率、断链距离或频谱结果。",
    )
    add_common_preface(doc)
    add_terms_two(doc)

    add_heading(doc, "二、相关技术背景（背景技术），与本发明最相近似的现有实现方案（现有技术）", 2)
    add_heading(doc, "2.1 技术领域", 3)
    add_para(doc, "本发明属于电力无人机巡检航线评估、DEM/DSM 地形分析、无线通信链路风险评估、GIS 风险图层和山区巡检通信保障技术领域。")
    add_heading(doc, "2.2 背景技术", 3)
    add_para(doc, "山区电力线路巡检场景下，无人机航点高度、通信源位置、山体或地表遮挡、Fresnel 区净空、飞行距离、近塔电磁暴露候选和气象条件可能共同影响通信链路可靠性。仅依靠二维距离、三维视觉底图或单一链路预算，难以在航点级别解释风险来源。")
    add_para(doc, "在一种实施方式中，系统可读取任意来源生成的标准化航线对象，包括航点经纬度、高度、航线标识、杆塔空间对象和通信源候选对象。该实施方式不要求必须依赖前一份交底书所述的数据标准化系统；只要输入对象满足字段语义和坐标要求，即可进入本发明的联合评估流程。")
    add_heading(doc, "2.3 现有技术一及其缺点", 3)
    add_para(doc, "现有技术一通常在三维地图或地形底图上展示航线和地形，由人工观察判断是否存在遮挡风险。该方案缺少 DEM/DSM 数据状态管理，也未将 KML 高度、航线设定高度、地形高程、展示高度和真实 AGL 区分处理，容易把视觉地形效果误认为工程高程结论。")
    add_heading(doc, "2.4 现有技术二及其缺点", 3)
    add_para(doc, "现有技术二通常按通信源到航点的距离或公知 FSPL 公式进行链路预算估算。该方案忽略沿剖面地形采样、LoS/NLoS 状态、Fresnel 净空、近塔 EMI 暴露候选和 GIS 风险图层输出之间的组合关系。若通信设备参数缺失，仍可能输出不具备依据的覆盖结论。")
    add_heading(doc, "2.5 核心区别特征", 3)
    add_numbered(doc, [
        "以航点级对象统一处理航线、杆塔、通信源和 DEM/DSM 地形数据。",
        "区分 KML 高度、航线设定高度、地形高程、杆塔高度、展示高度和真实 AGL，避免高度语义混用。",
        "DEM/DSM 不可用时输出 DEM not provided 或 terrain pending，不计算真实 AGL、真实遮挡或工程覆盖结论。",
        "在通信源到航点的剖面采样链路上联合计算 LoS、可配置比例的第一 Fresnel 区净空、距离链路压力和近塔 EMI 暴露候选。",
        "输出航点级或航段级风险标签、GIS 风险图层、剖面证据图和待补充数据清单。"
    ])

    add_heading(doc, "三、本发明技术方案的详细阐述（技术方案）", 2)
    add_heading(doc, "3.1 本发明所要解决的技术问题（发明目的）", 3)
    add_para(doc, "本发明解决的技术问题在于：将航点、杆塔、通信源和 DEM/DSM 地形统一到同一航点级评估对象中；在地形或通信参数缺失时输出明确状态边界；在地形数据可用时沿通信链路剖面进行 LoS 和 Fresnel 净空判定；在通信参数不足时输出 distance-only 而非覆盖结论；并将评估结果转化为航点级或航段级风险标签及 GIS 风险图层。")
    add_heading(doc, "3.2 系统组成", 3)
    add_numbered(doc, [
        "标准化航线输入模块，用于读取任意来源生成的航点级航线对象。",
        "通信源管理模块，用于读取基站、机巢、中继或地面站候选点，并记录通信参数的可用状态。",
        "高度语义统一模块，用于区分 KML 高度、航线设定高度、地形高程、杆塔高度、展示高度和真实 AGL。",
        "DEM/DSM 状态管理模块，用于在地形数据缺失时返回 DEM not provided 或 terrain pending，在地形数据可用时提供高程采样。",
        "链路剖面采样模块，用于在通信源到航点之间生成剖面采样点。",
        "LoS 与 Fresnel 净空模块，用于计算 h_LoS、F1 和净空余量 C，并输出 LoS clear、Fresnel risk 或 NLoS blocked 等状态。",
        "距离链路压力模块，用于在坐标可用时输出距离压力；在设备参数不足时输出 distance-only。",
        "EMI 候选暴露模块，用于基于航点与杆塔或电力设备的空间邻近关系生成近塔电磁暴露候选风险，不输出实测电磁场强。",
        "风险融合与图层输出模块，用于生成航点级或航段级风险标签、GIS 风险图层、剖面证据图和待补充数据清单。"
    ])

    add_heading(doc, "3.3 方法流程", 3)
    add_numbered(doc, [
        "步骤A1，读取标准化航线对象、杆塔对象和通信源对象。",
        "步骤A2，检查 DEM/DSM 数据源和通信参数完整性；缺失参数标记为 parameter missing 或待工程标定。",
        "步骤A3，统一高度语义；仅当 DEM/DSM 可用时计算真实 AGL。",
        "步骤A4，对每个航点和通信源生成剖面采样点。",
        "步骤A5，计算链路线高度 h_LoS，并读取或判断地形高程 h_DEM。",
        "步骤A6，在频率参数可用时计算第一 Fresnel 区半径 F1，并根据可配置净空比例 k 计算净空余量 C。",
        "步骤A7，根据状态机输出 terrain pending、distance-only、LoS clear、Fresnel risk、NLoS blocked 或 parameter missing。",
        "步骤A8，计算近塔 EMI 暴露候选风险，仅作为空间邻近和作业阶段提示。",
        "步骤A9，融合多个分量生成航点级或航段级风险标签。",
        "步骤A10，输出 GIS 风险图层、剖面证据图和待补充数据清单。"
    ])

    add_heading(doc, "3.4 算法化表达", 3)
    add_para(doc, "设通信源 S 的水平位置为 (lon_s,lat_s)，航点 P 的水平位置为 (lon_p,lat_p)。将 S 到 P 的链路按 M 个采样点离散，u_m=m/(M-1)，m=0,...,M-1。采样点 Q_m 可由大地线插值、局部投影线性插值或其他空间插值方法得到，M 为可配置参数。")
    add_para(doc, "若 DEM/DSM 可用，则通信源绝对高度 H_s 可表示为通信源地形高程加通信源天线高度，航点绝对高度 H_p 可依据已确认的高度语义计算。链路线高度为 h_LoS(u_m)=H_s+u_m*(H_p-H_s)。若 DEM/DSM 不可用，则不计算真实 H_s、H_p 和真实 AGL。")
    add_para(doc, "当频率 f 可用时，波长 lambda=c/f。设采样点到通信源距离为 d1(u_m)，到航点距离为 d2(u_m)，则第一 Fresnel 区半径可表示为 F1(u_m)=sqrt(lambda*d1(u_m)*d2(u_m)/(d1(u_m)+d2(u_m)))。净空余量 C(u_m)=h_LoS(u_m)-h_DEM(u_m)-k*F1(u_m)，其中 k 为可配置净空比例，常用值可由工程要求标定，本文不写死。")
    add_code_block(doc, [
        "for each waypoint P and communication source S:",
        "    if DEM/DSM is unavailable:",
        "        terrain_state = terrain_pending or DEM not provided",
        "        if coordinates are available: output distance-only",
        "        do not compute real AGL or real blockage",
        "    elif required communication parameter is missing:",
        "        compute terrain profile when possible",
        "        output parameter missing for link-budget-dependent items",
        "    else:",
        "        generate profile samples Q_m",
        "        compute h_LoS(Q_m), F1(Q_m), C(Q_m)",
        "        if any terrain height >= h_LoS:",
        "            status = NLoS blocked",
        "        elif min(C) < 0:",
        "            status = Fresnel risk",
        "        else:",
        "            status = LoS clear",
        "    compute near-tower EMI candidate from spatial proximity only",
        "    fuse status, distance pressure, EMI candidate, weather and data completeness",
    ])
    add_para(doc, "风险融合规则可采用规则表、专家权重或经实测数据校准的模型实现。若状态为 NLoS blocked，则该航点或航段至少标记为高通信敏感候选；若状态为 Fresnel risk，则标记为中高通信敏感候选；若 DEM/DSM 缺失，则标记为 terrain pending；若通信参数不足，则标记为 parameter missing 或 distance-only。所有权重、阈值和风险等级边界均为可配置或待工程标定。")
    add_para(doc, "EMI 分量仅表示近塔电磁暴露候选风险。其输入可以包括航点到最近杆塔或电力设备的距离、是否处于近塔悬停或绕塔巡检阶段、天气修正项和设备类型标记。该分量不输出实测电磁场强、频谱功率或设备抗扰度结论，除非后续接入经核验的实测数据。")

    add_heading(doc, "3.5 状态机说明", 3)
    add_table(doc, [
        ["状态", "触发条件", "输出边界"],
        ["terrain_pending / DEM not provided", "DEM/DSM 不存在、不可读或未完成采样", "不计算真实 AGL、真实遮挡或工程覆盖结论"],
        ["distance-only", "通信源和航点坐标可用，但地形或设备参数不足", "仅输出距离压力或阈值提示，不输出覆盖质量"],
        ["parameter missing", "频段、天线、功率、接收灵敏度、带宽或权重等必要参数缺失", "相关链路预算或融合分量标记待补充"],
        ["LoS clear", "地形高程未遮断链路线且净空余量满足配置条件", "表示模型计算状态，不替代实测验收"],
        ["Fresnel risk", "未形成 NLoS，但净空余量 C 小于配置阈值", "输出菲涅尔净空不足风险候选"],
        ["NLoS blocked", "至少一个采样点地形或表面高程高于链路线高度", "输出非视距遮挡风险候选"],
    ], [2100, 3750, 3510], header=True)

    add_heading(doc, "3.6 实施例", 3)
    add_para(doc, "在一种实施方式中，系统读取标准化航线对象和基站候选点，选择距航线起点最近或由人工指定的通信源。若 DEM/DSM 尚未接入，系统在状态字段中输出 DEM not provided，并仅输出航点至通信源距离和 needs_los_analysis 标记，不计算真实 AGL 或地形遮挡。")
    add_para(doc, "在另一种实施方式中，DEM/DSM 已接入且频段、天线高度等参数可用。系统对通信源到每个航点之间的剖面进行采样，计算 h_LoS、F1 和 C，生成 LoS clear、Fresnel risk 或 NLoS blocked 标签，并将结果聚合为航段风险色带、剖面证据图和待补充数据清单。")
    add_figure_descriptions_two(doc)

    add_heading(doc, "四、技术方案产生的效果（本发明技术方案带来的有益效果）", 2)
    add_numbered(doc, [
        "在航点级对象上统一处理航线、杆塔、通信源和 DEM/DSM 地形数据，提高风险解释粒度。",
        "通过高度语义统一，避免把 KML 高度、展示高度或航线设定高度误认为真实 AGL。",
        "通过 DEM/DSM 状态管理，在地形数据缺失时明确阻断真实遮挡和 AGL 结论。",
        "通过剖面采样联合计算 LoS、Fresnel 净空和距离压力，形成比单一距离阈值更可解释的通信敏感区识别流程。",
        "通过近塔 EMI 暴露候选分量，将电力巡检近塔作业特征纳入风险提示，但不伪造实测电磁场强。",
        "通过 GIS 风险图层、剖面证据图和待补充数据清单，便于后续外业复核和工程标定。"
    ])

    add_heading(doc, "五、可替代方案", 2)
    add_para(doc, "DEM/DSM 可采用 GeoTIFF、HGT、ASC、受控高程服务或其他可追溯地形数据实现；剖面采样可采用大地线插值或局部投影插值；通信传播模型可采用 FSPL、绕射模型、经验模型、射线追踪或经实测校准的模型；风险融合可采用规则表、专家权重、层次分析或经标签数据训练的模型。上述替代方案均应保留数据源、参数来源和缺失状态。")

    add_heading(doc, "六、技术关键点和欲保护点", 2)
    add_numbered(doc, [
        "保护以航点级对象统一处理航线、杆塔、通信源和 DEM/DSM 地形数据的联合评估流程。",
        "保护高度语义统一机制，区分 KML 高度、航线设定高度、地形高程、杆塔高度、展示高度和真实 AGL。",
        "保护 DEM/DSM 不可用时输出 DEM not provided 或 terrain pending 且不计算真实 AGL、真实遮挡或工程覆盖结论的状态边界管理。",
        "保护通信源到航点之间的剖面采样、h_LoS、F1 和 C 的组合计算流程。",
        "保护 distance-only、parameter missing、LoS clear、Fresnel risk 和 NLoS blocked 等状态机输出。",
        "保护将近塔 EMI 空间邻近性作为候选暴露风险而非实测电磁结论的处理机制。",
        "保护多因子风险标签、GIS 风险图层、剖面证据图和待补充数据清单的组合输出。"
    ])

    add_heading(doc, "七、权利要求草案", 2)
    add_heading(doc, "7.1 独立权利要求草案", 3)
    add_numbered(doc, [
        "一种地形净空与通信链路风险联合评估系统，包括：标准化航线输入模块、通信源管理模块、高度语义统一模块、DEM/DSM 状态管理模块、链路剖面采样模块、LoS 与 Fresnel 净空模块、距离链路压力模块、EMI 候选暴露模块、风险融合与图层输出模块；其中，系统以航点级对象统一处理航线、杆塔、通信源和 DEM/DSM 地形数据，并在 DEM/DSM 不可用时输出 DEM not provided 或 terrain pending，禁止计算真实 AGL、真实遮挡或工程覆盖结论。",
        "一种地形净空与通信链路风险联合评估方法，包括：读取航点级航线对象、杆塔对象和通信源对象；检查 DEM/DSM 和通信参数状态；统一高度语义；生成通信源到航点之间的剖面采样点；在地形数据可用时计算 h_LoS、F1 和 C；根据状态机输出 terrain pending、distance-only、parameter missing、LoS clear、Fresnel risk 或 NLoS blocked；生成近塔 EMI 暴露候选风险；融合生成航点级或航段级风险标签、GIS 风险图层、剖面证据图和待补充数据清单。",
        "一种电子设备或计算机可读存储介质，其上存储有计算机程序，所述计算机程序被处理器执行时实现上述地形净空与通信链路风险联合评估方法。"
    ])
    add_heading(doc, "7.2 从属权利要求要点", 3)
    add_bullets(doc, [
        "DEM/DSM 状态管理进一步包括在地形数据缺失、不可读或采样不足时输出 terrain pending。",
        "高度语义统一进一步包括分别保存 KML 高度、航线设定高度、地形高程、杆塔高度、展示高度和真实 AGL。",
        "链路剖面采样进一步包括按照可配置采样数量在通信源与航点之间生成 Q_m。",
        "LoS 判定进一步包括比较 h_LoS 与 h_DEM 的相对关系。",
        "Fresnel 净空进一步包括采用可配置比例 k 计算 C=h_LoS-h_DEM-kF1。",
        "distance-only 输出进一步包括在坐标可用但地形或设备参数不足时仅输出距离压力。",
        "EMI 候选暴露进一步包括基于航点与杆塔或电力设备的空间邻近关系生成候选风险。",
        "多因子风险融合进一步包括将状态机结果、距离压力、EMI 候选、天气和数据完整性合成为风险标签。",
        "GIS 图层输出进一步包括风险航段、风险点、通信走廊、剖面证据图和待补充数据清单。"
    ])

    add_prior_art_search(doc, [
        "DEM 视距分析",
        "Fresnel 区通信链路评估",
        "无人机空地信道",
        "GIS 风险图层",
        "电力巡检通信保障",
        "电磁暴露评估"
    ])
    add_heading(doc, "九、参考文献（如专利/论文/标准等）", 2)
    add_para(doc, "公开专利、论文、标准和 DOI 信息均待人工检索核验；本文不伪造检索结论。")
    save_doc(doc, OUT_DIR / "专利交底书2_地形净空与通信链路风险联合评估系统及方法.docx")


def save_doc(doc: Document, path: Path) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(path)


def main() -> None:
    build_disclosure_one()
    build_disclosure_two()


if __name__ == "__main__":
    main()
